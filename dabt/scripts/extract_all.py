#!/usr/bin/env python3
"""
Extract all questions from Chapter Tests, Kristen Topic Tests, and Kristen Mini Exams
into standardized CSVs matching the DABT database schema.
"""

import csv
import os
import re
from docx import Document

# ─── HELPERS ─────────────────────────────────────────────────────────

def clean_text(txt):
    """Clean up text: normalize whitespace, remove extra newlines."""
    txt = re.sub(r'\s+', ' ', txt).strip()
    return txt

def extract_questions_from_exam(filepath):
    """
    Extract questions and options from an exam file.
    Returns list of (question_number, question_text, {letter: option_text}, ordered_letters)
    """
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    
    questions = []
    i = 0
    while i < len(paragraphs):
        txt = paragraphs[i]
        if not txt:
            i += 1
            continue
        
        # Detect question: starts with digit(s) followed by '.' or ')'
        m = re.match(r'^(\d+)[).]\s*(.*)', txt)
        if m:
            qnum = m.group(1)
            qtext = m.group(2).strip()
            i += 1
            
            options = {}
            opt_order = []
            
            while i < len(paragraphs):
                nxt = paragraphs[i]
                if not nxt:
                    i += 1
                    continue
                
                # Next question
                if re.match(r'^\d+[).]', nxt):
                    break
                
                # Option line: letter followed by '.' or ')' or ' '
                m2 = re.match(r'^([A-H])[).\s]\s*(.*)', nxt)
                if m2:
                    letter = m2.group(1)
                    rest = m2.group(2).strip()
                    # Sometimes options have extra text (in answer keys). Strip annotations.
                    # Extract just the option text (before any colon or explanation markers)
                    options[letter] = rest
                    if letter not in opt_order:
                        opt_order.append(letter)
                
                i += 1
            
            if qtext:  # Only add if there's actual question text
                questions.append((qnum, qtext, options, opt_order))
    
    return questions


def extract_answers_from_star_key(filepath):
    """
    Extract answer explanations from a docx with * markers.
    Returns dict: question_number -> answer_text
    """
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    
    # First find all questions and their options
    questions_info = {}  # qnum -> {'text': ..., 'options': {}, 'order': []}
    
    i = 0
    while i < len(paragraphs):
        txt = paragraphs[i]
        if not txt:
            i += 1
            continue
        
        m = re.match(r'^(\d+)[).]\s*(.*)', txt)
        if m:
            qnum = m.group(1)
            qtext = m.group(2).strip()
            i += 1
            
            options = {}
            opt_order = []
            
            while i < len(paragraphs):
                nxt = paragraphs[i]
                if not nxt:
                    i += 1
                    continue
                if re.match(r'^\d+[).]', nxt):
                    break
                m2 = re.match(r'^([A-H])[).\s]\s*(.*)', nxt)
                if m2:
                    letter = m2.group(1)
                    rest = m2.group(2).strip()
                    options[letter] = rest
                    if letter not in opt_order:
                        opt_order.append(letter)
                i += 1
            
            if qtext and qnum not in questions_info:
                questions_info[qnum] = {'text': qtext, 'options': options, 'order': opt_order}
    
    # Now find * paragraphs and associate them with questions
    answers = {}  # qnum -> list of answer texts
    current_qnum = None
    
    for i, txt in enumerate(paragraphs):
        if not txt:
            continue
        
        # Check if this is a question
        m = re.match(r'^(\d+)[).]', txt)
        if m:
            current_qnum = m.group(1)
            continue
        
        # Check for * marker
        if txt.startswith('*'):
            answer_text = txt[1:].strip()
            if current_qnum:
                if current_qnum not in answers:
                    answers[current_qnum] = []
                answers[current_qnum].append(answer_text)
    
    return questions_info, answers


def match_answer_letter(options, opt_order, answer_texts):
    """
    Given options dict and answer explanation text(s), determine the correct answer letter.
    Uses keyword matching.
    """
    if not answer_texts or not options:
        return None, ""
    
    # Join all answer texts
    combined = ' '.join(answer_texts).lower()
    
    best_letter = None
    best_score = -1
    scores = {}
    
    for letter in opt_order:
        opt_text = options.get(letter, '').lower().strip()
        if not opt_text:
            continue
        
        score = 0
        
        # Full option text appears in answer
        if len(opt_text) > 3 and opt_text in combined:
            score += 100
        
        # Option text at start of any answer paragraph
        for ans in answer_texts:
            ans_lower = ans.lower()
            if ans_lower.startswith(opt_text[:min(len(opt_text), 80)]):
                score += 200
                break
        
        # Key words match (words > 4 chars)
        words = [w for w in re.findall(r'\b[a-zA-Z]{4,}\b', opt_text) if w not in ('with', 'that', 'this', 'from', 'than', 'they', 'have', 'been', 'were', 'their', 'which', 'what', 'will', 'when', 'does', 'into', 'also', 'some', 'more', 'most', 'over', 'about', 'after', 'other', 'would', 'could', 'should', 'these', 'those', 'there', 'where', 'being')]
        if words:
            match_count = sum(1 for w in words if w in combined)
            score += match_count * 5
        
        # For compound options like "all of the above", "A and B", "C and D"
        # Check if mentioned options are all correct
        compound_match = re.match(r'^([A-H])\s*and\s*([A-H])$', opt_text, re.IGNORECASE)
        compound_all = re.match(r'^(all|none)\s*of\s*the\s*above$', opt_text, re.IGNORECASE)
        
        if compound_match:
            l1, l2 = compound_match.group(1).upper(), compound_match.group(2).upper()
            # Check if both options' texts appear in the answer
            t1 = options.get(l1, '').lower()
            t2 = options.get(l2, '').lower()
            if t1 and t2:
                if len(t1) > 3 and t1 in combined:
                    score += 50
                if len(t2) > 3 and t2 in combined:
                    score += 50
                if len(t1) > 3 and t1 in combined and len(t2) > 3 and t2 in combined:
                    score += 100  # Bonus if both match
        
        if compound_all:
            # Check if most individual options appear in the answer
            individual_options = [options.get(l, '').lower() for l in opt_order if l != letter]
            matches = sum(1 for t in individual_options if len(t) > 3 and t in combined)
            if matches >= len(individual_options) * 0.5:
                score += 80
        
        scores[letter] = score
        if score > best_score:
            best_score = score
            best_letter = letter
    
    if best_letter and best_score > 0:
        return best_letter, options.get(best_letter, '')
    
    return None, ""


def determine_inline_answer(options, opt_order):
    """
    For inline-annotation answer keys (Mini Exams), determine the correct answer.
    Look for patterns: "(incorrect)" markers, explanation text length, etc.
    Returns (letter, text) or (None, "")
    """
    if not options:
        return None, ""
    
    # Strategy 1: Look for options NOT marked as incorrect/wrong/false
    # In many inline answer keys, wrong answers are marked with (incorrect) or similar
    has_incorrect_markers = any(
        'incorrect' in t.lower() or '(wrong)' in t.lower() or 'false' in t.lower() 
        for t in options.values()
    )
    
    if has_incorrect_markers:
        # The correct answer is the one WITHOUT incorrect markers
        for letter in opt_order:
            txt = options.get(letter, '')
            if not any(marker in txt.lower() for marker in ['incorrect', '(wrong)', 'false ']):
                # Also check if it has reasonable text (not too short)
                if len(txt) > 5:
                    return letter, re.split(r'[:]', txt, 1)[0].strip() if ':' in txt else txt
    
    # Strategy 2: Check for options with explanation text (longer)
    # In some files, the correct answer has extensive explanation
    # But this can also match wrong answers that explain why they're wrong
    # So we use this as a secondary heuristic
    
    # Strategy 3: Check the length distribution
    lengths = [(letter, len(options.get(letter, ''))) for letter in opt_order]
    # Sort by length descending
    lengths.sort(key=lambda x: -x[1])
    
    # If one option is significantly longer than others, it might be the correct one
    if len(lengths) >= 3:
        longest_len = lengths[0][1]
        second_len = lengths[1][1]
        if longest_len > second_len * 3 and longest_len > 100:
            letter = lengths[0][0]
            txt = options.get(letter, '')
            # But check: if the text says "incorrect" or "wrong", it's a wrong explanation
            if not any(marker in txt.lower() for marker in ['incorrect', '(wrong)']):
                return letter, re.split(r'[:]', txt, 1)[0].strip() if ':' in txt else txt
    
    return None, ""


# ─── CHAPTER MAPPING ─────────────────────────────────────────────────

CHAPTER_TOPIC_MAP = {
    'Chapter 2': 'General Principles & Concepts',
    'Chapter 3': 'Mechanisms of Toxicity',
    'Chapter 5': 'General Toxicology',
    'Chapter 6A': 'Biotransformation / Metabolism',
    'Chapter 6B': 'Biotransformation / Metabolism',
    'Chapter 7': 'Toxicokinetics / ADME',
    'Chapter 8': 'Carcinogenesis & Mutagenesis',
    'Chapter 9': 'Genotoxicity / DNA Damage',
    'Chapters 10': 'Reproductive & Developmental Toxicity',
    'Chapter 11': 'Hematology & Blood Toxicity',
    'Chapter 12A': 'Immunotoxicology / Allergy',
    'Chapter 12B': 'Immunotoxicology / Allergy',
    'Chapter 13': 'Liver / Hepatotoxicity',
    'Chapter 14': 'Kidney / Nephrotoxicity',
    'Chapter 15': 'Lung / Pulmonary Toxicity',
    'Chapter 16': 'Nervous System / Neurotoxicity',
    'Chapter 17': 'Eye / Ocular Toxicity',
    'Chapter 18': 'Cardiovascular Toxicity',
    'Chapter 19': 'Skin / Dermatotoxicity',
    'Chapter 21': 'Endocrine Toxicology',
    'Chapter 23': 'Metals & Metalloids',
    'Chapter 24': 'Solvents & Hydrocarbons',
    'Chapters 26': 'Plant Toxins',
    'Chapters 28': 'Air Pollution & Particulates',
    'Chapter 30': 'Food Additives, Cosmetics & GRAS',
    'Chapter 20': 'Reproductive & Developmental Toxicity',
}

def get_chapter_topic(filename):
    """Map chapter filename to topic."""
    for key, topic in CHAPTER_TOPIC_MAP.items():
        if key in filename:
            return topic
    return 'General Principles & Concepts'

TOPIC_MAP_FILENAME = {
    'air pollution': 'Air Pollution & Particulates',
    'carcinogenicity': 'Carcinogenesis & Mutagenesis',
    'ecotox': 'General Toxicology',
    'endocrine': 'Endocrine Toxicology',
    'forensic': 'Drugs & Therapeutics – Toxicology',
    'genotoxicity': 'Genotoxicity / DNA Damage',
    'oculartox': 'Eye / Ocular Toxicity',
    'pesticides': 'Pesticides – Insecticides',
    'repro': 'Reproductive & Developmental Toxicity',
    'respiratory': 'Lung / Pulmonary Toxicity',
    'risk assess': 'Risk Assessment & Regulatory',
    'skin': 'Skin / Dermatotoxicity',
    'solvents': 'Solvents & Hydrocarbons',
}

def get_topic_topic(filename):
    fname_lower = filename.lower()
    for key, topic in TOPIC_MAP_FILENAME.items():
        if key in fname_lower:
            return topic
    return 'General Principles & Concepts'


def get_source_exam_name(filename, bank='chapter'):
    """Get a short readable source exam name."""
    if bank == 'chapter':
        # Extract chapter number and name
        m = re.match(r'Chapters?\s+(\d+[A-Za-z]?(?:\s+and\s+\d+[A-Za-z]?)?)\s*[-–]\s*(.+)\.docx', filename)
        if m:
            num = m.group(1).strip()
            topic = m.group(2).strip()
            return f"Chapter {num} ({topic[:40]})"
        return filename.replace('.docx', '')
    
    elif bank == 'topic':
        # e.g., "Air pollution examination 20 July 2017.docx"
        m = re.match(r'^([A-Za-z\s-]+?)\s*(?:examination|exam)\s.*', filename)
        if m:
            topic = m.group(1).strip()
            return f"{topic.title()} Topic"
        return filename.replace('.docx', '')
    
    elif bank == 'mini':
        # e.g., "Mini-ABT exam 09 June 2017.docx"
        m = re.match(r'^Mini-ABT\s*(?:exam|examination)\s*(.*)\.docx', filename)
        if m:
            date = m.group(1).strip()
            return f"Mini-ABT {date}"
        return filename.replace('.docx', '')
    
    return filename.replace('.docx', '')


def get_chapter_source_name(filename):
    """Get a nice source exam name for chapter tests."""
    f = filename.replace('.docx', '')
    # Extract chapter info
    m = re.match(r'(Chapters?\s+\d+[A-Za-z]?(?:\s+and\s+\d+[A-Za-z]?)?)\s*[-–]\s*(.*)', f)
    if m:
        num_part = m.group(1).strip()
        topic_part = m.group(2).strip()
        # Shorten
        short_topics = {
            'Toxic Responses of the Blood': 'Blood',
            'Toxic Responses of the Immune System': 'Immune System',
            'Toxic Responses of the Liver': 'Liver',
            'Toxic Responses of the Kidney': 'Kidney',
            'Toxic Responses of the Respiratory System': 'Respiratory',
            'Toxic Responses of the Nervous System': 'Nervous System',
            'Toxic Responses of the Ocular and Visual System': 'Ocular',
            'Toxic Responses of the Heart and Vascular System': 'Cardiovascular',
            'Toxic Responses of the Skin': 'Skin',
            'Toxic Responses of the Endocrine System': 'Endocrine',
            'Toxic Effects of Metals': 'Metals',
            'Toxic Effects of Solvents and Vapors': 'Solvents',
            'General Principles': 'General Principles',
            'Mechanisms of Toxicity': 'Mechanisms',
            'Absorption Distribution and Excretion of Toxicants': 'ADE',
            'Biotransformation of Xenobiotics': 'Biotransformation',
            'Toxicokinetics': 'Toxicokinetics',
            'Chemical Carcinogens': 'Carcinogens',
            'Genetic Toxicology': 'Genetic Toxicology',
            'Developmental and Reproductive Tox': 'Dev/Repro',
            'Animal Venoms and Plants': 'Venoms/Plants',
            'Air Pollution and Ecotoxicology': 'Air/Ecotox',
            'Food Toxicology': 'Food Toxicology',
        }
        for long, short in short_topics.items():
            if long in topic_part:
                return f"Chapter {num_part} {short}"
        return f"Chapter {num_part} {topic_part[:30]}"
    return f


def get_topic_source_name(filename):
    f = filename.replace('.docx', '')
    m = re.match(r'^([A-Za-z\s/-]+?)\s*(?:examination|exam)\s.*', f)
    if m:
        topic = m.group(1).strip()
        return f"{topic.title()} Topic"
    return f


def get_mini_source_name(filename):
    f = filename.replace('.docx', '')
    for prefix in ['Mini-ABT exam for ', 'Mini-ABT examination for ', 
                   'Mini-ABT exam ', 'Mini-ABT examination ', 'Mini-ABT ']:
        if prefix in f:
            f = f.replace(prefix, '', 1)
            break
    f = f.strip()
    return f"Mini-ABT {f}"


# ─── BANK A: CHAPTER TESTS ──────────────────────────────────────────

def process_chapter_tests():
    """Extract questions from Chapter Tests."""
    test_dir = "/root/dabt-curated/Chapter_Tests/Tests/"
    answer_dir = "/root/dabt-curated/Chapter_Tests/Tests_with_Answers/"
    
    all_rows = []
    
    # Map test files to answer key files
    test_files = sorted([f for f in os.listdir(test_dir) if f.endswith('.docx')])
    
    for test_file in test_files:
        print(f"Processing chapter test: {test_file}")
        
        # Find answer key
        base_name = test_file.replace('.docx', '')
        answer_file = None
        for af in os.listdir(answer_dir):
            if af.endswith('.docx') and base_name in af:
                answer_file = af
                break
        
        if not answer_file:
            print(f"  WARNING: No answer key found for {test_file}")
            continue
        
        # Extract questions from exam file
        qs = extract_questions_from_exam(os.path.join(test_dir, test_file))
        
        # Extract answers from answer key
        q_info, answers = extract_answers_from_star_key(os.path.join(answer_dir, answer_file))
        
        # Source info
        source_name = get_chapter_source_name(test_file)
        topic = get_chapter_topic(test_file)
        source_file = test_file
        
        for qnum, qtext, options, opt_order in qs:
            # Get answer explanation
            answer_texts = answers.get(qnum, [])
            
            # Determine correct answer
            correct_letter, correct_text = match_answer_letter(options, opt_order, answer_texts)
            
            # Build row
            row = {
                'ID': '',
                'Source Exam': source_name,
                'Question #': qnum,
                'Question': qtext,
                'A': options.get('A', ''),
                'B': options.get('B', ''),
                'C': options.get('C', ''),
                'D': options.get('D', ''),
                'E': options.get('E', ''),
                'F': options.get('F', ''),
                'G': options.get('G', ''),
                'H': options.get('H', ''),
                'Correct Answer': correct_letter or '',
                'Correct Answer Text': correct_text,
                'Explanation': ' '.join(answer_texts) if answer_texts else '',
                'Topic (Primary)': topic,
                'All Topics': topic,
                'Source File': source_file,
            }
            all_rows.append(row)
        
        print(f"  -> {len(qs)} questions extracted")
    
    return all_rows


# ─── BANK B: TOPIC TESTS ────────────────────────────────────────────

def process_topic_tests():
    """Extract questions from Kristen Topic Tests."""
    topic_dir = "/root/dabt-curated/Practice_Tests_by_Topic/Kristen_Topic_Tests/"
    
    all_rows = []
    test_files = sorted([f for f in os.listdir(topic_dir) if f.endswith('.docx')
                         and not any(p in f.lower() for p in ['answer', 'explanation', 'explaination', 'answers'])])
    answer_files_list = sorted([f for f in os.listdir(topic_dir) if f.endswith('.docx')
                                and any(p in f.lower() for p in ['answer', 'explanation', 'explaination', 'answers'])])
    
    for test_file in test_files:
        print(f"Processing topic test: {test_file}")
        
        # Find matching answer key
        # Strategy: extract base topic name and find the answer key
        base_key = test_file.lower()
        # Remove words like "examination", "exam", dates
        base_key = re.sub(r'\s+(?:examination|exam)\s+.*', '', base_key)
        
        answer_file = None
        for af in answer_files_list:
            af_lower = af.lower()
            if base_key in af_lower:
                answer_file = af
                break
        
        if not answer_file:
            # Try matching by first part of filename
            m = re.match(r'^([A-Za-z\s-]+?)\s*(?:examination|exam)\s', test_file)
            if m:
                topic_part = m.group(1).strip().lower()
                for af in answer_files_list:
                    if topic_part in af.lower():
                        answer_file = af
                        break
        
        # Extract questions from exam file
        qs = extract_questions_from_exam(os.path.join(topic_dir, test_file))
        
        # Extract answers from answer key
        q_info = {}
        answers = {}
        if answer_file:
            q_info, answers = extract_answers_from_star_key(os.path.join(topic_dir, answer_file))
        else:
            print(f"  WARNING: No answer key found for {test_file}")
        
        source_name = get_topic_source_name(test_file)
        topic = get_topic_topic(test_file)
        source_file = test_file
        
        for qnum, qtext, options, opt_order in qs:
            answer_texts = answers.get(qnum, [])
            
            correct_letter, correct_text = match_answer_letter(options, opt_order, answer_texts)
            
            row = {
                'ID': '',
                'Source Exam': source_name,
                'Question #': qnum,
                'Question': qtext,
                'A': options.get('A', ''),
                'B': options.get('B', ''),
                'C': options.get('C', ''),
                'D': options.get('D', ''),
                'E': options.get('E', ''),
                'F': options.get('F', ''),
                'G': options.get('G', ''),
                'H': options.get('H', ''),
                'Correct Answer': correct_letter or '',
                'Correct Answer Text': correct_text,
                'Explanation': ' '.join(answer_texts) if answer_texts else '',
                'Topic (Primary)': topic,
                'All Topics': topic,
                'Source File': source_file,
            }
            all_rows.append(row)
        
        print(f"  -> {len(qs)} questions extracted")
    
    return all_rows


# ─── BANK C: MINI EXAMS ─────────────────────────────────────────────

def process_mini_exams():
    """Extract questions from Kristen Mini Exams."""
    mini_dir = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams/"
    
    all_rows = []
    
    # Exam files (without answers)
    exclude_patterns = ['answer', 'explanation', 'explaination', 'answers']
    test_files = sorted([f for f in os.listdir(mini_dir) if f.endswith('.docx')
                         and not any(p in f.lower() for p in exclude_patterns)])
    answer_files_list = sorted([f for f in os.listdir(mini_dir) if f.endswith('.docx')
                                and any(p in f.lower() for p in exclude_patterns)])
    
    for test_file in test_files:
        print(f"Processing mini exam: {test_file}")
        
        # Find matching answer key
        # Extract date pattern from test file name
        date_pattern = re.search(r'(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})', test_file, re.IGNORECASE)
        date_pattern2 = re.search(r'(\d{1,2}\s+[A-Za-z]+\.?\s+\d{4})', test_file)
        
        answer_file = None
        if date_pattern or date_pattern2:
            date_str = (date_pattern or date_pattern2).group(1)
            # Normalize date for matching
            date_normalized = date_str.lower().replace('.', '').strip()
            
            for af in answer_files_list:
                af_lower = af.lower().replace('.docx', '')
                af_lower = af_lower.replace('.', '')
                if date_normalized in af_lower:
                    answer_file = af
                    break
        else:
            # Fallback: try part number matching (A/B)
            for af in answer_files_list:
                af_lower = af.lower()
                # Check if key parts match
                test_parts = set(test_file.lower().replace('.docx', '').split())
                af_parts = set(af_lower.replace('.docx', '').split())
                common = test_parts & af_parts
                if len(common) >= 3 and 'with' not in test_parts:
                    answer_file = af
                    break
        
        if not answer_file:
            print(f"  WARNING: No answer key found for {test_file}")
        
        # Extract questions from exam file
        qs = extract_questions_from_exam(os.path.join(mini_dir, test_file))
        
        # Source info
        source_name = get_mini_source_name(test_file)
        source_file = test_file
        topic = 'General Principles & Concepts'  # Mini exams are comprehensive
        
        # Process answers
        if answer_file:
            answer_path = os.path.join(mini_dir, answer_file)
            
            # Check if answer key uses * markers or inline annotations
            doc = Document(answer_path)
            paras = [p.text.strip() for p in doc.paragraphs]
            star_count = sum(1 for p in paras if p.startswith('*'))
            
            if star_count > 5:
                # Uses * markers - use star-based extraction
                q_info, answers = extract_answers_from_star_key(answer_path)
                
                for qnum, qtext, options, opt_order in qs:
                    answer_texts = answers.get(qnum, [])
                    correct_letter, correct_text = match_answer_letter(options, opt_order, answer_texts)
                    
                    row = {
                        'ID': '',
                        'Source Exam': source_name,
                        'Question #': qnum,
                        'Question': qtext,
                        'A': options.get('A', ''),
                        'B': options.get('B', ''),
                        'C': options.get('C', ''),
                        'D': options.get('D', ''),
                        'E': options.get('E', ''),
                        'F': options.get('F', ''),
                        'G': options.get('G', ''),
                        'H': options.get('H', ''),
                        'Correct Answer': correct_letter or '',
                        'Correct Answer Text': correct_text,
                        'Explanation': ' '.join(answer_texts) if answer_texts else '',
                        'Topic (Primary)': topic,
                        'All Topics': 'Comprehensive',
                        'Source File': source_file,
                    }
                    all_rows.append(row)
            else:
                # Uses inline annotations - extract options from answer key directly
                # Parse the answer key file to find which options have explanations
                ans_qs = extract_questions_from_exam(answer_path)
                
                # Build mapping from question number to options with annotations
                ans_map = {}
                for aqnum, aqtext, aoptions, aopt_order in ans_qs:
                    ans_map[aqnum] = (aoptions, aopt_order)
                
                for qnum, qtext, options, opt_order in qs:
                    ans_opts, ans_order = ans_map.get(qnum, ({}, []))
                    
                    if ans_opts:
                        # Use the inline-annotated options to determine correct answer
                        correct_letter, correct_text = determine_inline_answer(ans_opts, ans_order)
                        
                        # The explanation is whatever extra text is in the annotated option
                        explanation = ''
                        if correct_letter and correct_letter in ans_opts:
                            txt = ans_opts[correct_letter]
                            # Extract explanation (text after the first colon or just the extra part)
                            parts = txt.split(':', 1)
                            if len(parts) > 1 and len(parts[1].strip()) > 20:
                                explanation = parts[1].strip()
                    else:
                        correct_letter, correct_text = None, ""
                        explanation = ''
                    
                    row = {
                        'ID': '',
                        'Source Exam': source_name,
                        'Question #': qnum,
                        'Question': qtext,
                        'A': options.get('A', ''),
                        'B': options.get('B', ''),
                        'C': options.get('C', ''),
                        'D': options.get('D', ''),
                        'E': options.get('E', ''),
                        'F': options.get('F', ''),
                        'G': options.get('G', ''),
                        'H': options.get('H', ''),
                        'Correct Answer': correct_letter or '',
                        'Correct Answer Text': correct_text or '',
                        'Explanation': explanation,
                        'Topic (Primary)': topic,
                        'All Topics': 'Comprehensive',
                        'Source File': source_file,
                    }
                    all_rows.append(row)
        else:
            # No answer key: just extract questions
            for qnum, qtext, options, opt_order in qs:
                row = {
                    'ID': '',
                    'Source Exam': source_name,
                    'Question #': qnum,
                    'Question': qtext,
                    'A': options.get('A', ''),
                    'B': options.get('B', ''),
                    'C': options.get('C', ''),
                    'D': options.get('D', ''),
                    'E': options.get('E', ''),
                    'F': options.get('F', ''),
                    'G': options.get('G', ''),
                    'H': options.get('H', ''),
                    'Correct Answer': '',
                    'Correct Answer Text': '',
                    'Explanation': '',
                    'Topic (Primary)': topic,
                    'All Topics': 'Comprehensive',
                    'Source File': source_file,
                }
                all_rows.append(row)
        
        print(f"  -> {len(qs)} questions extracted")
    
    return all_rows


# ─── CSV WRITING ─────────────────────────────────────────────────────

FIELDNAMES = ['ID', 'Source Exam', 'Question #', 'Question', 'A', 'B', 'C', 'D',
              'E', 'F', 'G', 'H', 'Correct Answer', 'Correct Answer Text',
              'Explanation', 'Topic (Primary)', 'All Topics', 'Source File']

def write_csv(filepath, rows, start_id_num):
    """Write rows to CSV with sequential IDs starting from start_id_num."""
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=FIELDNAMES)
        writer.writeheader()
        
        for i, row in enumerate(rows):
            row['ID'] = f"DABT-{start_id_num + i:04d}"
            writer.writerow(row)
    
    return start_id_num + len(rows)


# ─── MAIN ────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("=" * 60)
    print("BANK A: Chapter Tests")
    print("=" * 60)
    chapter_rows = process_chapter_tests()
    print(f"\nTotal chapter test questions: {len(chapter_rows)}\n")
    
    print("=" * 60)
    print("BANK B: Kristen Topic Tests")
    print("=" * 60)
    topic_rows = process_topic_tests()
    print(f"\nTotal topic test questions: {len(topic_rows)}\n")
    
    print("=" * 60)
    print("BANK C: Kristen Mini Exams")
    print("=" * 60)
    mini_rows = process_mini_exams()
    print(f"\nTotal mini exam questions: {len(mini_rows)}\n")
    
    # Write CSVs
    print("=" * 60)
    print("Writing CSVs...")
    
    next_id = 447  # Start at DABT-0447
    next_id = write_csv('/tmp/dabt_extract_chapter.csv', chapter_rows, next_id)
    print(f"Chapter CSV written, next ID: {next_id}")
    
    next_id = write_csv('/tmp/dabt_extract_topic.csv', topic_rows, next_id)
    print(f"Topic CSV written, next ID: {next_id}")
    
    next_id = write_csv('/tmp/dabt_extract_mini.csv', mini_rows, next_id)
    print(f"Mini CSV written, next ID: {next_id}")
    
    total = len(chapter_rows) + len(topic_rows) + len(mini_rows)
    print(f"\nTotal questions extracted: {total}")
    print(f"Final ID range: DABT-0447 to DABT-{447 + total - 1:04d}")
