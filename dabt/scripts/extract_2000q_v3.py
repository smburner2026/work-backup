#!/usr/bin/env python3
"""
Extract ALL ~1,937 questions from 2000Q Question Bank docx files into standardized CSV.
Version 3: Properly handles options on separate lines within the same paragraph.
"""

import os, re, csv, sys
from collections import OrderedDict, defaultdict
from docx import Document

BASE_DIR = "/root/dabt-curated/2000Q_Question_Bank"
ANSWER_KEY_CSV = os.path.join(BASE_DIR, "2000Q_ANSWER_KEY.csv")
OUTPUT_CSV = "/tmp/dabt_extract_2000q.csv"

DOCX_FILES = sorted([
    f for f in os.listdir(BASE_DIR)
    if f.endswith('.docx') and f.startswith('2000Q')
])

def source_exam_name(filename):
    name = filename.replace('2000Q ', '').replace('.docx', '')
    name = name.replace('_DUPLICATE_', '')
    mapping = {
        'Ch1 General': 'Ch1 General',
        'Ch3 Mechanisms of Toxicity': 'Ch3 Mechanisms',
        'Ch5 ADE': 'Ch5 ADE',
        'Ch6 Metabolism': 'Ch6 Metabolism',
        'Ch7 TK': 'Ch7 TK',
        'Ch8 Cardiovascular Tox': 'Ch8 Cardiovascular',
        'Ch11 Blood': 'Ch11 Blood',
        'Ch12 Immunology': 'Ch12 Immunology',
        'Ch13 Hepatic Tox': 'Ch13 Hepatic',
        'Ch14 Renal Tox': 'Ch14 Renal',
        'Ch15 RESPIRATORY TOXICOLOGY': 'Ch15 Respiratory',
        'Ch16 NEUROTOXICOLOGY': 'Ch16 Neurotoxicology',
        'Ch17 OCULAR TOXICOLOGY': 'Ch17 Ocular',
        'Ch18 Cardiovascular': 'Ch18 Cardiovascular',
        'Ch19 Dermal Tox': 'Ch19 Dermal',
        'Ch20 REPRODUCTIVE TOXICOLOGY': 'Ch20 Reproductive',
        'Ch21 ENDOCRINE TOXICOLOGY': 'Ch21 Endocrine',
        'Ch22 PESTICIDES': 'Ch22 Pesticides',
        'Ch23 METAL TOXICOLOGY': 'Ch23 Metal',
        'Ch24 CHEMICAL AND SOLVENT TOXICOLOGY': 'Ch24 Chemical & Solvent',
        'Ch25 RADIATION TOXICOLOGY': 'Ch25 Radiation',
        'Ch27 TOXICOLOGY OF PLANTS': 'Ch27 Plant Toxins',
        'Ch28 AIR AND WATER POLLUTION': 'Ch28 Air & Water',
        'Ch29 ENVIRONMENTAL TOXICOLOGY': 'Ch29 Environmental',
        'Ch30 FOOD TOXICOLOGY': 'Ch30 Food',
        'Ch31 ANALYTICAL TOXICOLOGY': 'Ch31 Analytical',
        'Ch33 OCCUPATIONAL TOXICOLOGY': 'Ch33 Occupational',
        'TOXICOLOGIC DISASTERS': 'Toxicologic Disasters',
        'EPIDEMIOLOGY': 'Epidemiology',
        'MEDICAL TOXICOLOGY': 'Medical Toxicology',
        'DRUG ABUSE TOXICOLOGY': 'Drug Abuse',
        'ANTIDOTES Matching Test': 'Antidotes',
        'ANIMAL TOXICOLOGY': 'Animal Toxicology',
        'ALCOHOL TOXICOLOGY': 'Alcohol Toxicology',
        'Ch11 HEMATOLOGIC TOX': 'Ch11 Hematologic',
    }
    for key, val in mapping.items():
        if key in name:
            return val
    return name.strip()

TOPIC_MAP = {
    'Ch1 General': 'General Principles & Concepts',
    'Ch3 Mechanisms': 'Mechanisms of Toxicity',
    'Ch5 ADE': 'General Toxicology',
    'Ch6 Metabolism': 'Biotransformation / Metabolism',
    'Ch7 TK': 'Toxicokinetics / ADME',
    'Ch8 Cardiovascular': 'Cardiovascular Toxicity',
    'Ch11 Blood': 'Hematology & Blood Toxicity',
    'Ch11 Hematologic': 'Hematology & Blood Toxicity',
    'Ch12 Immunology': 'Immunotoxicology / Allergy',
    'Ch13 Hepatic': 'Liver / Hepatotoxicity',
    'Ch14 Renal': 'Kidney / Nephrotoxicity',
    'Ch15 Respiratory': 'Lung / Pulmonary Toxicity',
    'Ch16 Neurotoxicology': 'Nervous System / Neurotoxicity',
    'Ch17 Ocular': 'Eye / Ocular Toxicity',
    'Ch18 Cardiovascular': 'Cardiovascular Toxicity',
    'Ch19 Dermal': 'Skin / Dermatotoxicity',
    'Ch20 Reproductive': 'Reproductive & Developmental Toxicity',
    'Ch21 Endocrine': 'Endocrine Toxicology',
    'Ch22 Pesticides': 'Pesticides – Insecticides',
    'Ch23 Metal': 'Metals & Metalloids',
    'Ch24 Chemical & Solvent': 'Solvents & Hydrocarbons',
    'Ch25 Radiation': 'Radiation / UV / Ionizing',
    'Ch27 Plant Toxins': 'Plant Toxins',
    'Ch28 Air & Water': 'Air Pollution & Particulates',
    'Ch29 Environmental': 'General Toxicology',
    'Ch30 Food': 'Food Additives, Cosmetics & GRAS',
    'Ch31 Analytical': 'General Toxicology',
    'Ch33 Occupational': 'General Toxicology',
    'Alcohol Toxicology': 'Alcohols & Methanol/Ethanol',
    'Animal Toxicology': 'General Toxicology',
    'Antidotes': 'Drugs & Therapeutics – Toxicology',
    'Drug Abuse': 'Drugs & Therapeutics – Toxicology',
    'Epidemiology': 'Risk Assessment & Regulatory',
    'Medical Toxicology': 'Drugs & Therapeutics – Toxicology',
    'Toxicologic Disasters': 'General Toxicology',
}

def get_topic(source_name):
    for key, topic in TOPIC_MAP.items():
        if key.lower() in source_name.lower():
            return topic
    return 'General Toxicology'

# Load answer key
print("Loading answer key...")
answer_key = {}
with open(ANSWER_KEY_CSV, 'r') as f:
    reader = csv.DictReader(f)
    for row in reader:
        qid = int(row['QuestionID'])
        answer_key[qid] = row.get('AnswerLetter', '').strip()
print(f"  Loaded {len(answer_key)} answer letters")

# Scan docx for actual question numbers
print("\nScanning docx files for question numbers...")
file_question_ids = defaultdict(set)

for fname in DOCX_FILES:
    filepath = os.path.join(BASE_DIR, fname)
    if not os.path.exists(filepath):
        continue
    doc = Document(filepath)
    qids = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Require period after number
            m = re.match(r'^(\d{1,4})\.\s+(.*)', line)
            if not m:
                # Also handle tab after number
                m = re.match(r'^(\d{1,4})\.\t(.*)', line)
            if m:
                qnum = int(m.group(1))
                if 1 <= qnum <= 1999:
                    rest = m.group(2).strip()
                    # Skip obvious chapter headers (all-caps title)
                    if rest and re.match(r'^[A-Z\s]{15,}$', rest):
                        continue
                    qids.add(qnum)
    if qids:
        file_question_ids[fname] = qids
        print(f"  {fname}: Q{min(qids)}-Q{max(qids)} ({len(qids)} questions)")

# Build final mapping
qid_to_file = {}
for qid in range(1, 2000):
    containing_files = []
    for fname, qids in file_question_ids.items():
        if qid in qids:
            containing_files.append(fname)
    if not containing_files:
        continue
    non_dup = [f for f in containing_files if not f.startswith('_DUPLICATE')]
    if non_dup:
        qid_to_file[qid] = non_dup[0]
    else:
        qid_to_file[qid] = containing_files[0]

print(f"\n  Mapped {len(qid_to_file)} questions to files")

# Organize by file
file_qid_list = defaultdict(list)
for qid in sorted(qid_to_file.keys()):
    fname = qid_to_file[qid]
    file_qid_list[fname].append(qid)

# Check for gaps
mapped_set = set(qid_to_file.keys())
all_expected = set(range(1, 2000))
missing = sorted(all_expected - mapped_set)
if missing:
    groups = []
    start = missing[0]
    end = missing[0]
    for m in missing[1:]:
        if m == end + 1:
            end = m
        else:
            groups.append((start, end))
            start = m
            end = m
    groups.append((start, end))
    print(f"\n  *** MISSING {len(missing)} QUESTIONS ***")
    for s, e in groups:
        if s == e:
            print(f"    Q{s}")
        else:
            print(f"    Q{s}-Q{e}")

# ============================================================
# EXTRACTION FUNCTION (Version 3)
# ============================================================
def parse_options_line(text):
    """
    Parse a line containing multiple options like:
    "A. option text B. option text C. option text D. option text"
    or "A) option B) option"
    Returns dict of {letter: text}
    """
    options = OrderedDict()
    # Pattern to find option markers
    # A letter A-H followed by period or paren, then text until next letter-period/paren or end
    # Use a lookahead to handle overlapping matches
    pattern = r'([A-Z])[\.\)]\s*((?:(?!\s+[A-Z][\.\)]).)+)'
    matches = list(re.finditer(pattern, text))
    
    for i, m in enumerate(matches):
        letter = m.group(1)
        opt_text = m.group(2).strip()
        # Clean up: remove trailing whitespace and stray chars
        opt_text = re.sub(r'\s+', ' ', opt_text)
        if opt_text and len(opt_text) > 0:
            # Make sure this isn't just a reference marker like "(I)" 
            if not re.match(r'^\([A-Z]+\)$', opt_text):
                options[letter] = opt_text
    
    return options


def extract_questions(filepath, known_qids):
    """
    Extract questions and answer options from docx.
    Process each paragraph line by line.
    """
    doc = Document(filepath)
    results = {}
    
    # Get all paragraph texts, splitting by newlines
    lines = []
    for para in doc.paragraphs:
        text = para.text
        for subline in text.split('\n'):
            subline = subline.strip()
            if subline:
                lines.append(subline)
    
    # Find answer key section start
    answer_start = len(lines)
    for i, line in enumerate(lines):
        if ('ANSWER' in line.upper() and 
            ('REFERENCE' in line.upper() or 'CHAPTER' in line.upper() or '(' in line)):
            # The answer section begins
            answer_start = i - 1  # Include the header for safety
            break
    
    if answer_start < 0:
        answer_start = 0
    
    content_lines = lines[:answer_start]
    
    # Remove trailing reference text (lines starting with references like "I. Author...")
    # Find the start of the references section
    ref_start = len(content_lines)
    for i in range(len(content_lines)):
        line = content_lines[i]
        if re.match(r'^[IVX]+\.\s+[A-Z]', line) and len(line) > 30:
            ref_start = i
            break
    
    content_lines = content_lines[:ref_start]
    
    # Remove empty trailing lines
    while content_lines and not content_lines[-1].strip():
        content_lines.pop()
    
    # For matching tests and files with inline options, do a simple pre-scan
    # to identify which lines start a new question
    
    # State machine
    current_qid = None
    current_question_parts = []
    current_options = OrderedDict()
    in_options_block = False
    
    qnum_re = re.compile(r'^(\d{1,4})\.\s+(.*)')
    opt_re = re.compile(r'^([A-Z])[\.\)]\s*(.*)')
    
    def save_current():
        nonlocal current_qid, current_question_parts, current_options, in_options_block
        if current_qid is not None and current_qid in known_qids:
            qtext = ' '.join(current_question_parts).strip()
            qtext = re.sub(r'\s+', ' ', qtext)
            qtext = qtext.replace('\t', ' ').strip()
            results[current_qid] = {
                'question': qtext,
                'options': dict(current_options),
            }
        current_qid = None
        current_question_parts = []
        current_options = OrderedDict()
        in_options_block = False
    
    for line in content_lines:
        # Try matching test inline format first: "N. term \t LETTER. definition"
        # or "N. term  LETTER. definition" (2+ spaces between term and letter)
        match_test = re.match(r'^(\d{1,4})\.\s+(.+?)\s{2,}([A-Z])[\.\)]\s*(.*)', line)
        if not match_test:
            match_test = re.match(r'^(\d{1,4})\.\s+(.+?)\t+([A-Z])[\.\)]?\s*(.*)', line)
        
        if match_test:
            qnum = int(match_test.group(1))
            term = match_test.group(2).strip()
            opt_letter = match_test.group(3).strip()
            opt_text = match_test.group(4).strip()
            
            if qnum in known_qids:
                save_current()
                current_qid = qnum
                current_question_parts = [term]
                current_options = OrderedDict()
                current_options[opt_letter] = opt_text
                save_current()
            continue
        
        # Check for question start
        m_q = qnum_re.match(line)
        if m_q:
            qnum = int(m_q.group(1))
            rest = m_q.group(2).strip()
            
            if qnum not in known_qids:
                save_current()
                continue
            
            # Save previous question
            save_current()
            
            current_qid = qnum
            current_question_parts = [rest]
            current_options = OrderedDict()
            in_options_block = False
            
            # Also need to look for inline options with A-Z
            if len(rest) > 5:
                opts = parse_options_line(rest)
                if opts:
                    # Find where the first option starts
                    first_opt_pattern = re.search(r'\s+([A-Z])[\.\)]\s+', rest)
                    if first_opt_pattern:
                        qtext = rest[:first_opt_pattern.start()].strip()
                        current_question_parts = [qtext]
                        current_options = opts
                        in_options_block = True
                        save_current()
                        continue
            
            # If we got here, just the question text - options will come later
            continue
        
        # Check for option lines (A. text or A) text)
        m_opt = opt_re.match(line)
        if m_opt:
            letter = m_opt.group(1)
            opt_text = m_opt.group(2).strip()
            
            if current_qid is not None and current_qid in known_qids:
                # Parse the rest of the line for multiple options
                # Example: "A. opt1 B. opt2 C. opt3 D. opt4"
                all_opts = parse_options_line(line)
                if all_opts:
                    current_options.update(all_opts)
                    in_options_block = True
                elif opt_text and len(opt_text) > 0:
                    # Single option (might be a continuation from previous line)
                    current_options[letter] = opt_text
                    in_options_block = True
            continue
        
        # If we're in options mode and the line doesn't start with a letter,
        # it might be a continuation of the last option
        if in_options_block and current_qid is not None:
            # Could be a continuation of the option text
            # Check if it starts with text (continuation of last option)
            if current_options:
                last_letter = list(current_options.keys())[-1]
                current_options[last_letter] = current_options[last_letter] + ' ' + line
            continue
        
        # If we're in a question, this is continuation text
        if current_qid is not None and current_qid in known_qids:
            current_question_parts.append(line)
            in_options_block = False
            continue
    
    save_current()
    return results


# ============================================================
# MAIN PROCESSING
# ============================================================
print("\n\nExtracting questions from docx files...")
next_id_num = 447
output_rows = []
extraction_stats = defaultdict(lambda: {'total': 0, 'extracted': 0, 'no_opts': 0, 'no_text': 0})

for fname in sorted(file_qid_list.keys()):
    filepath = os.path.join(BASE_DIR, fname)
    expected_qids = file_qid_list[fname]
    
    print(f"\n  {fname} ({len(expected_qids)} questions, Q{expected_qids[0]}-Q{expected_qids[-1]})")
    
    extracted = extract_questions(filepath, set(expected_qids))
    
    source_name = source_exam_name(fname)
    topic = get_topic(source_name)
    
    stats = extraction_stats[fname]
    stats['total'] = len(expected_qids)
    stats['extracted'] = len(extracted)
    
    sorted_qids = sorted(expected_qids)
    qnum_in_file = 0
    
    for global_qid in sorted_qids:
        qnum_in_file += 1
        qinfo = extracted.get(global_qid, {})
        question_text = qinfo.get('question', '')
        options = qinfo.get('options', {})
        
        if not question_text:
            stats['no_text'] += 1
        
        if not options:
            stats['no_opts'] += 1
        
        correct_letter = answer_key.get(global_qid, '')
        correct_text = options.get(correct_letter, '')
        
        dabt_id = f"DABT-{next_id_num:04d}"
        next_id_num += 1
        
        row = OrderedDict()
        row['ID'] = dabt_id
        row['Source Exam'] = source_name
        row['Question #'] = qnum_in_file
        row['Question'] = question_text
        row['A'] = options.get('A', '')
        row['B'] = options.get('B', '')
        row['C'] = options.get('C', '')
        row['D'] = options.get('D', '')
        row['E'] = options.get('E', '')
        row['F'] = options.get('F', '')
        row['G'] = options.get('G', '')
        row['H'] = options.get('H', '')
        row['Correct Answer'] = correct_letter
        row['Correct Answer Text'] = correct_text
        row['Explanation'] = ''
        row['Topic (Primary)'] = topic
        row['All Topics'] = topic
        row['Source File'] = fname
        
        output_rows.append(row)
    
    if stats['no_opts']:
        print(f"    No options for {stats['no_opts']} questions")
    if stats['no_text']:
        print(f"    No text for {stats['no_text']} questions")

# Write CSV
print(f"\n\nWriting CSV to {OUTPUT_CSV}...")
fieldnames = list(output_rows[0].keys()) if output_rows else []
with open(OUTPUT_CSV, 'w', newline='') as f:
    writer = csv.DictWriter(f, fieldnames=fieldnames)
    writer.writeheader()
    for row in output_rows:
        writer.writerow(row)

print(f"  Written {len(output_rows)} questions")
print(f"  ID range: {output_rows[0]['ID']} to {output_rows[-1]['ID']}")

# Summary
print(f"\n{'='*70}")
print("EXTRACTION SUMMARY")
print(f"{'='*70}")
print(f"{'File':50s} {'Total':>5s} {'Extr.':>5s} {'NoOpt':>5s} {'NoTxt':>5s}")
print(f"{'-'*70}")
for fname in sorted(extraction_stats.keys()):
    s = extraction_stats[fname]
    print(f"{fname:50s} {s['total']:5d} {s['extracted']:5d} {s['no_opts']:5d} {s['no_text']:5d}")

total_found = sum(s['total'] for s in extraction_stats.values())
total_no_opts = sum(s['no_opts'] for s in extraction_stats.values())
total_no_text = sum(s['no_text'] for s in extraction_stats.values())
print(f"{'-'*70}")
print(f"{'TOTAL':50s} {total_found:5d} {total_found:5d} {total_no_opts:5d} {total_no_text:5d}")

print("\nDONE!")
