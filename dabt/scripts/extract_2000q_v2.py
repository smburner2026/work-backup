#!/usr/bin/env python3
"""
Extract ALL ~1,937 questions from 2000Q Question Bank docx files into standardized CSV
matching the DABT database schema.

Version 2: More robust question extraction using content-based file mapping.
"""

import os, re, csv, sys
from collections import OrderedDict, defaultdict
from docx import Document

BASE_DIR = "/root/dabt-curated/2000Q_Question_Bank"
ANSWER_KEY_CSV = os.path.join(BASE_DIR, "2000Q_ANSWER_KEY.csv")
OUTPUT_CSV = "/tmp/dabt_extract_2000q.csv"

# Include ALL 2000Q docx files (including _DUPLICATE_ since they may be the only copy)
DOCX_FILES = sorted([
    f for f in os.listdir(BASE_DIR)
    if f.endswith('.docx') and f.startswith('2000Q')
])

def source_exam_name(filename):
    """Derive a short human-readable source exam name from filename."""
    name = filename.replace('2000Q ', '').replace('.docx', '')
    name = name.replace('_DUPLICATE_', '')
    # Short, human-readable names
    mapping = {
        'Ch1 General': 'Ch1 General',
        'Ch3 Mechanisms of Toxicity': 'Ch3 Mechanisms',
        'Ch5 ADE': 'Ch5 ADE',
        'Ch6 Metabolism': 'Ch6 Metabolism',
        'Ch7 TK': 'Ch7 TK',
        'Ch8 Cardiovascular Tox': 'Ch8 Cardiovascular',
        'Ch11 Blood': 'Ch11 Blood',
        'Ch12 Immunology': 'Ch12 Immunology',
        'Ch13 Hepatic Tox': 'Ch13 Hepatic',
        'Ch14 Renal Tox': 'Ch14 Renal',
        'Ch15 RESPIRATORY TOXICOLOGY': 'Ch15 Respiratory',
        'Ch16 NEUROTOXICOLOGY': 'Ch16 Neurotoxicology',
        'Ch17 OCULAR TOXICOLOGY': 'Ch17 Ocular',
        'Ch18 Cardiovascular': 'Ch18 Cardiovascular',
        'Ch19 Dermal Tox': 'Ch19 Dermal',
        'Ch20 REPRODUCTIVE TOXICOLOGY': 'Ch20 Reproductive',
        'Ch21 ENDOCRINE TOXICOLOGY': 'Ch21 Endocrine',
        'Ch22 PESTICIDES': 'Ch22 Pesticides',
        'Ch23 METAL TOXICOLOGY': 'Ch23 Metal',
        'Ch24 CHEMICAL AND SOLVENT TOXICOLOGY': 'Ch24 Chemical & Solvent',
        'Ch25 RADIATION TOXICOLOGY': 'Ch25 Radiation',
        'Ch27 TOXICOLOGY OF PLANTS': 'Ch27 Plant Toxins',
        'Ch28 AIR AND WATER POLLUTION': 'Ch28 Air & Water',
        'Ch29 ENVIRONMENTAL TOXICOLOGY': 'Ch29 Environmental',
        'Ch30 FOOD TOXICOLOGY': 'Ch30 Food',
        'Ch31 ANALYTICAL TOXICOLOGY': 'Ch31 Analytical',
        'Ch33 OCCUPATIONAL TOXICOLOGY': 'Ch33 Occupational',
        'TOXICOLOGIC DISASTERS': 'Toxicologic Disasters',
        'EPIDEMIOLOGY': 'Epidemiology',
        'MEDICAL TOXICOLOGY': 'Medical Toxicology',
        'DRUG ABUSE TOXICOLOGY': 'Drug Abuse',
        'ANTIDOTES Matching Test': 'Antidotes',
        'ANIMAL TOXICOLOGY': 'Animal Toxicology',
        'ALCOHOL TOXICOLOGY': 'Alcohol Toxicology',
        'Ch11 HEMATOLOGIC TOX': 'Ch11 Hematologic',
    }
    for key, val in mapping.items():
        if key in name:
            return val
    return name.strip()

TOPIC_MAP = {
    'Ch1 General': 'General Principles & Concepts',
    'Ch3 Mechanisms': 'Mechanisms of Toxicity',
    'Ch5 ADE': 'General Toxicology',
    'Ch6 Metabolism': 'Biotransformation / Metabolism',
    'Ch7 TK': 'Toxicokinetics / ADME',
    'Ch8 Cardiovascular': 'Cardiovascular Toxicity',
    'Ch11 Blood': 'Hematology & Blood Toxicity',
    'Ch11 Hematologic': 'Hematology & Blood Toxicity',
    'Ch12 Immunology': 'Immunotoxicology / Allergy',
    'Ch13 Hepatic': 'Liver / Hepatotoxicity',
    'Ch14 Renal': 'Kidney / Nephrotoxicity',
    'Ch15 Respiratory': 'Lung / Pulmonary Toxicity',
    'Ch16 Neurotoxicology': 'Nervous System / Neurotoxicity',
    'Ch17 Ocular': 'Eye / Ocular Toxicity',
    'Ch18 Cardiovascular': 'Cardiovascular Toxicity',
    'Ch19 Dermal': 'Skin / Dermatotoxicity',
    'Ch20 Reproductive': 'Reproductive & Developmental Toxicity',
    'Ch21 Endocrine': 'Endocrine Toxicology',
    'Ch22 Pesticides': 'Pesticides – Insecticides',
    'Ch23 Metal': 'Metals & Metalloids',
    'Ch24 Chemical & Solvent': 'Solvents & Hydrocarbons',
    'Ch25 Radiation': 'Radiation / UV / Ionizing',
    'Ch27 Plant Toxins': 'Plant Toxins',
    'Ch28 Air & Water': 'Air Pollution & Particulates',
    'Ch29 Environmental': 'General Toxicology',
    'Ch30 Food': 'Food Additives, Cosmetics & GRAS',
    'Ch31 Analytical': 'General Toxicology',
    'Ch33 Occupational': 'General Toxicology',
    'Alcohol Toxicology': 'Alcohols & Methanol/Ethanol',
    'Animal Toxicology': 'General Toxicology',
    'Antidotes': 'Drugs & Therapeutics – Toxicology',
    'Drug Abuse': 'Drugs & Therapeutics – Toxicology',
    'Epidemiology': 'Risk Assessment & Regulatory',
    'Medical Toxicology': 'Drugs & Therapeutics – Toxicology',
    'Toxicologic Disasters': 'General Toxicology',
}

def get_topic(source_name):
    for key, topic in TOPIC_MAP.items():
        if key.lower() in source_name.lower():
            return topic
    return 'General Toxicology'

# Step 1: Load answer key
print("Loading answer key...")
answer_key = {}  # global_qid -> answer_letter
with open(ANSWER_KEY_CSV, 'r') as f:
    reader = csv.DictReader(f)
    for row in reader:
        qid = int(row['QuestionID'])
        answer_key[qid] = row.get('AnswerLetter', '').strip()
print(f"  Loaded {len(answer_key)} answer letters")

# Step 2: Scan each docx to find which global question IDs are present
print("\nScanning docx files for question numbers...")
file_question_ids = defaultdict(set)  # filename -> set of global qids

for fname in DOCX_FILES:
    filepath = os.path.join(BASE_DIR, fname)
    if not os.path.exists(filepath):
        continue
    
    doc = Document(filepath)
    qids = set()
    
    for para in doc.paragraphs:
        text = para.text.strip()
        # Must have a period after the number to be a real question,
        # or a tab. Chapter headers like "1 TOXICOLOGIC PRINCIPLES" have just space.
        m = re.match(r'^(\d{1,4})\.\s+(.*)', text)
        if not m:
            m = re.match(r'^(\d{1,4})\.\t(.*)', text)
        if m:
            qnum = int(m.group(1))
            if 1 <= qnum <= 1999:
                # Skip obvious chapter headers (all-caps or mostly-caps after number)
                rest = m.group(2).strip() if m.lastindex >= 2 else ''
                if rest and re.match(r'^[A-Z\s]{15,}$', rest):
                    continue  # Chapter header like "TOXICOLOGIC PRINCIPLES"
                qids.add(qnum)
    
    if qids:
        file_question_ids[fname] = qids
        print(f"  {fname}: Q{min(qids)}-Q{max(qids)} ({len(qids)} questions)")

# Step 3: Build final file-to-question mapping
# For each global QID, find which file contains it
# If multiple files contain the same QID (from _DUPLICATE_), prefer the non-duplicate
print("\nBuilding question-to-file mapping...")
qid_to_file = {}
for qid in range(1, 2000):
    containing_files = []
    for fname, qids in file_question_ids.items():
        if qid in qids:
            containing_files.append(fname)
    
    if not containing_files:
        continue
    
    # Prefer non-duplicate files
    non_dup = [f for f in containing_files if not f.startswith('_DUPLICATE')]
    if non_dup:
        qid_to_file[qid] = non_dup[0]  # Use first non-duplicate
    else:
        qid_to_file[qid] = containing_files[0]  # Use duplicate

print(f"  Mapped {len(qid_to_file)} questions to files")

# Step 4: Organize questions by file
file_qid_list = defaultdict(list)
for qid in sorted(qid_to_file.keys()):
    fname = qid_to_file[qid]
    file_qid_list[fname].append(qid)

# Report
total_mapped = sum(len(v) for v in file_qid_list.values())
print(f"\n  Total questions mapped: {total_mapped}")
print(f"  Total files: {len(file_qid_list)}")
for fname in sorted(file_qid_list.keys()):
    qids = file_qid_list[fname]
    print(f"    {fname}: Q{qids[0]}-Q{qids[-1]} ({len(qids)} questions)")

# Verify covered range
mapped_set = set(qid_to_file.keys())
all_expected = set(range(1, 2000))
missing = sorted(all_expected - mapped_set)
if missing:
    print(f"\n  *** MISSING QUESTIONS: {len(missing)} ***")
    # Group consecutive missing
    groups = []
    start = missing[0]
    end = missing[0]
    for m in missing[1:]:
        if m == end + 1:
            end = m
        else:
            groups.append((start, end))
            start = m
            end = m
    groups.append((start, end))
    for s, e in groups:
        if s == e:
            print(f"    Q{s}")
        else:
            print(f"    Q{s}-Q{e}")

# Step 5: Extract questions from each docx
def extract_questions_v3(filepath, known_qids):
    """
    Extract questions and answer options from docx.
    Handles:
    - Regular MCQs: "N. Question text" then "A. opt" "B. opt" etc.
    - Inline options: "N. Question A. opt B. opt C. opt D. opt"
    - Matching tests: "N. term \t LETTER. definition"
    - Multi-line question text
    - Answer key sections at the end
    """
    doc = Document(filepath)
    results = {}  # qid -> {question, options}
    
    # Get all paragraphs (keep empty ones for structure detection)
    all_paras = [para.text for para in doc.paragraphs]
    
    # Find answer key section start
    answer_start = len(all_paras)
    for i, text in enumerate(all_paras):
        stripped = text.strip()
        if ('ANSWER' in stripped.upper() and 
            ('REFERENCE' in stripped.upper() or 'CHAPTER' in stripped.upper() or '(' in stripped)):
            answer_start = i
            break
    
    # Process paragraphs before answer key
    content = all_paras[:answer_start]
    
    # Remove empty trailing paragraphs
    while content and not content[-1].strip():
        content.pop()
    
    # State machine
    current_qid = None
    current_question_parts = []
    current_options = OrderedDict()
    
    def save_current():
        nonlocal current_qid, current_question_parts, current_options
        if current_qid is not None and current_qid in known_qids:
            qtext = ' '.join(current_question_parts).strip()
            qtext = re.sub(r'\s+', ' ', qtext)
            # Clean up any remaining tab characters
            qtext = qtext.replace('\t', ' ')
            results[current_qid] = {
                'question': qtext,
                'options': dict(current_options),
            }
        current_qid = None
        current_question_parts = []
        current_options = OrderedDict()
    
    # Regex patterns - require period after number to distinguish from chapter headers
    qnum_re = re.compile(r'^(\d{1,4})\.\s+(.*)')
    opt_re = re.compile(r'^([A-H])[\.\)]\s*(.*)')
    
    # First pass: identify paragraphs that contain question starts
    # For matching tests, the pattern is "N. term LETTER. definition" on one line
    # For regular Qs, it's "N. text" then optionally "A. opt" etc.
    
    i = 0
    while i < len(content):
        text = content[i].strip()
        
        if not text:
            i += 1
            continue
        
        # Check for matching test format: "NUM. term \t LETTER. definition"
        match_inline = re.match(r'^(\d{1,4})\.\s+(.+?)\s{2,}([A-H])[\.\)]\s*(.*)', text)
        if not match_inline:
            match_inline = re.match(r'^(\d{1,4})\.\s+(.+?)\t+([A-H])[\.\)]?\s*(.*)', text)
        
        if match_inline:
            qnum = int(match_inline.group(1))
            term = match_inline.group(2).strip()
            opt_letter = match_inline.group(3).strip()
            opt_text = match_inline.group(4).strip()
            
            if qnum in known_qids:
                save_current()
                current_qid = qnum
                current_question_parts = [term]
                current_options = OrderedDict()
                current_options[opt_letter] = opt_text
                save_current()
            i += 1
            continue
        
        # Check for question start
        m_q = qnum_re.match(text)
        if m_q:
            qnum = int(m_q.group(1))
            rest = m_q.group(2).strip()
            
            if qnum not in known_qids:
                save_current()
                i += 1
                continue
            
            save_current()
            current_qid = qnum
            
            # Check if this line has inline options
            # Look for pattern like "text A. opt B. opt C. opt" after some text
            # First, find where the first option starts
            opt_start = None
            for match in re.finditer(r'\s+([A-H])[\.\)]\s+', rest):
                pos = match.start()
                # Check if the letter is really an option marker (followed by non-trivial text)
                letter = match.group(1)
                after_match = rest[match.end():]
                # Parse what follows
                next_opt = re.search(r'\s+([A-H])[\.\)]\s+', after_match)
                if next_opt:
                    opt_text_end = next_opt.start()
                    opt_text_match = after_match[:opt_text_end].strip()
                else:
                    opt_text_match = after_match.strip()
                
                if opt_text_match and len(opt_text_match) > 0:
                    opt_start = pos
                    break
            
            if opt_start is not None:
                # Question text is before the first option
                qtext = rest[:opt_start].strip()
                current_question_parts = [qtext]
                
                # Parse all inline options
                opt_part = rest[opt_start:].strip()
                # Split on letter patterns
                opts = re.findall(r'([A-H])[\.\)]\s*((?:(?!\s+[A-H][\.\)]\s).)+)', opt_part)
                for letter, opt_text in opts:
                    opt_text = re.sub(r'\s+', ' ', opt_text.strip())
                    if opt_text and letter not in current_options:
                        current_options[letter] = opt_text
                
                save_current()
                i += 1
                continue
            else:
                # Just a question number with text
                current_question_parts = [rest]
                i += 1
                continue
        
        # Check for answer option lines
        m_opt = opt_re.match(text)
        if m_opt and current_qid is not None:
            letter = m_opt.group(1)
            opt_text = m_opt.group(2).strip()
            if opt_text and letter not in current_options:
                current_options[letter] = opt_text
            i += 1
            continue
        
        # If we're in a question and this isn't an option, it's continuation text
        if current_qid is not None:
            # Skip lines that look like section headers
            if re.match(r'^\d+\s+[A-Z\s]+$', text) and len(text) < 60:
                # Could be a chapter header like "12 CARDIOVASCULAR TOXICOLOGY"
                # But also could be matching test content
                # Check if it looks like a chapter/section title
                if re.match(r'^\d+\s+[A-Z\s]{10,}', text):
                    i += 1
                    continue
            current_question_parts.append(text)
            i += 1
            continue
        
        i += 1
    
    # Save last question
    save_current()
    
    return results


# Step 6: Process each file and generate output
print("\n\nExtracting questions from docx files...")
next_id_num = 447
output_rows = []
extraction_stats = defaultdict(lambda: {'total': 0, 'extracted': 0, 'no_opts': 0, 'no_text': 0})

for fname in sorted(file_qid_list.keys()):
    filepath = os.path.join(BASE_DIR, fname)
    expected_qids = file_qid_list[fname]
    
    print(f"\n  {fname} ({len(expected_qids)} questions)")
    
    extracted = extract_questions_v3(filepath, set(expected_qids))
    
    source_name = source_exam_name(fname)
    topic = get_topic(source_name)
    
    stats = extraction_stats[fname]
    stats['total'] = len(expected_qids)
    stats['extracted'] = len(extracted)
    
    # Sort by qid
    sorted_qids = sorted(expected_qids)
    qnum_in_file = 0
    
    for global_qid in sorted_qids:
        qnum_in_file += 1
        qinfo = extracted.get(global_qid, {})
        question_text = qinfo.get('question', '')
        options = qinfo.get('options', {})
        
        if not question_text:
            stats['no_text'] += 1
        
        if not options:
            stats['no_opts'] += 1
        
        correct_letter = answer_key.get(global_qid, '')
        correct_text = options.get(correct_letter, '')
        
        dabt_id = f"DABT-{next_id_num:04d}"
        next_id_num += 1
        
        row = OrderedDict()
        row['ID'] = dabt_id
        row['Source Exam'] = source_name
        row['Question #'] = qnum_in_file
        row['Question'] = question_text
        row['A'] = options.get('A', '')
        row['B'] = options.get('B', '')
        row['C'] = options.get('C', '')
        row['D'] = options.get('D', '')
        row['E'] = options.get('E', '')
        row['F'] = options.get('F', '')
        row['G'] = options.get('G', '')
        row['H'] = options.get('H', '')
        row['Correct Answer'] = correct_letter
        row['Correct Answer Text'] = correct_text
        row['Explanation'] = ''
        row['Topic (Primary)'] = topic
        row['All Topics'] = topic
        row['Source File'] = fname
        
        output_rows.append(row)
    
    if stats['no_opts']:
        print(f"    No options for {stats['no_opts']} questions")
    if stats['no_text']:
        print(f"    No text for {stats['no_text']} questions")

# Step 7: Write CSV
print(f"\n\nWriting CSV to {OUTPUT_CSV}...")
fieldnames = list(output_rows[0].keys()) if output_rows else []
with open(OUTPUT_CSV, 'w', newline='') as f:
    writer = csv.DictWriter(f, fieldnames=fieldnames)
    writer.writeheader()
    for row in output_rows:
        writer.writerow(row)

print(f"  Written {len(output_rows)} questions")
print(f"  ID range: {output_rows[0]['ID']} to {output_rows[-1]['ID']}")

# Summary by file
print(f"\n{'='*70}")
print("EXTRACTION SUMMARY")
print(f"{'='*70}")
print(f"{'File':50s} {'Total':>5s} {'Extr.':>5s} {'NoOpt':>5s} {'NoTxt':>5s}")
print(f"{'-'*70}")
for fname in sorted(extraction_stats.keys()):
    s = extraction_stats[fname]
    print(f"{fname:50s} {s['total']:5d} {s['extracted']:5d} {s['no_opts']:5d} {s['no_text']:5d}")

total_found = sum(s['total'] for s in extraction_stats.values())
total_extracted = sum(s['extracted'] for s in extraction_stats.values())
total_no_opts = sum(s['no_opts'] for s in extraction_stats.values())
total_no_text = sum(s['no_text'] for s in extraction_stats.values())
print(f"{'-'*70}")
print(f"{'TOTAL':50s} {total_found:5d} {total_extracted:5d} {total_no_opts:5d} {total_no_text:5d}")

print("\nDONE!")
