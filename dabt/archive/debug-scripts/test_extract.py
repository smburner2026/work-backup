#!/usr/bin/env python3
"""Examine a docx to count questions and extract answers."""
import sys, os, re
from docx import Document

def count_questions(doc_path):
    doc = Document(doc_path)
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        # Match "N." at start of line where N is a number
        if re.match(r'^\d+\.', text):
            count += 1
    return count

def extract_answers(doc_path):
    """Extract question_number -> answer_letter from a 'with answers' docx."""
    doc = Document(doc_path)
    answers = {}
    current_q = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this is a question number line
        m = re.match(r'^(\d+)\.\s+(.*)', text)
        if m:
            current_q = int(m.group(1))
            # Don't set answer yet, it's the question text
        
        # Check for bold answer option (e.g., "A. text" or "B. text" or "A. ...")
        bold_texts = []
        for run in para.runs:
            if run.bold and run.text.strip():
                bold_texts.append(run.text.strip())
        
        if bold_texts:
            full_bold = ''.join(bold_texts).strip()
            # Check if bold text starts with a letter option like "A." or "B." etc.
            opt_match = re.match(r'^([A-H])\.?\s*', full_bold)
            if opt_match and current_q is not None:
                letter = opt_match.group(1)
                answers[current_q] = letter
    
    return answers

# Test with one file
exam_dir = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"

# Test the "with answers" version
with_answers = os.path.join(exam_dir, "Mini-ABT exam with answers 26 August 2017.docx")
print(f"=== {os.path.basename(with_answers)} ===")
answers = extract_answers(with_answers)
print(f"Extracted {len(answers)} answers")
for qnum in sorted(answers.keys())[:5]:
    print(f"  Q{qnum}: {answers[qnum]}")
print("  ...")
for qnum in sorted(answers.keys())[-5:]:
    print(f"  Q{qnum}: {answers[qnum]}")
print(f"  Full: {dict(sorted(answers.items()))}")

# Also count questions in the non-answers version
non_answers = os.path.join(exam_dir, "Mini-ABT exam 26 August 2017.docx")
qcount = count_questions(non_answers)
print(f"\nNon-answers file has ~{qcount} questions")
