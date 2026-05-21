#!/usr/bin/env python3
"""Debug problematic docx files that extracted 0 answers."""
from docx import Document

files_to_check = [
    "Mini-ABT exam with answers 05 May 2017.docx",
    "Mini-ABT exam with answers 09 June 2017.docx",
    "Mini-ABT exam with answers 21 July 2017.docx",
    "Mini-ABT exam with answers for 11 August 2017.docx",
    "Mini-ABT exam with answers 12 May 2017.docx",
]

exam_dir = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"

for fname in files_to_check:
    fpath = f"{exam_dir}/{fname}"
    print(f"\n=== {fname} ===")
    doc = Document(fpath)
    
    # Print all paragraphs
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if not text:
            continue
        bold_runs = [r.text.strip() for r in para.runs if r.bold]
        bold_indicator = " [BOLD: " + "|".join(bold_runs) + "]" if bold_runs else ""
        print(f"  PARA {i}: {text[:120]}{bold_indicator}")
    
    # Check for tables
    for ti, table in enumerate(doc.tables):
        print(f"  TABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows[:5]):
            for ci, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"    Cell[{ri},{ci}]: {cell_text[:80]}")
