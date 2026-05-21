#!/usr/bin/env python3
"""Deep debug of tricky docx files - show ALL bold text in context."""
from docx import Document

exam_dir = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"

files_to_check = [
    "Mini-ABT exam with answers 05 May 2017.docx",
    "Mini-ABT exam with answers 09 June 2017.docx",
    "Mini-ABT exam with answers 21 July 2017.docx",
    "Mini-ABT exam with answers for 11 August 2017.docx",
    "Mini-ABT exam with answers for 15 Sept. 2017.docx",
    "Mini-ABT exam with answers for 22 Sept. 2017 (corrected.ver.2).docx",
    "Mini-ABT examination with answers 02 June 2017.docx",
    "Mini-ABT exam with answers 03 September 2017.docx",
    "Mini-ABT exam with answers 26 August 2017.docx",
]

for fname in files_to_check:
    fpath = f"{exam_dir}/{fname}"
    doc = Document(fpath)
    
    print(f"\n{'='*80}")
    print(f"FILE: {fname}")
    print(f"{'='*80}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        bold_runs = [r.text for r in para.runs if r.bold]
        if bold_runs:
            bold_text = ''.join(bold_runs).strip()
            print(f"  BOLD: '{bold_text[:100]}'")
            print(f"  FULL: '{text[:120]}'")
            print()

    print(f"  [Total paragraphs with bold: {sum(1 for p in doc.paragraphs if any(r.bold for r in p.runs))}]")
