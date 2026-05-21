#!/usr/bin/env python3
"""Debug: check what answers v4 extracted for specific files and qnums."""
import os, re
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"

# Use the same function from v4
def extract_answers_v4(doc_path):
    doc = Document(doc_path)
    answers = {}
    current_q = None
    options_since_q = {}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        m = re.match(r'^(\d+)[\.\)]\s+', text)
        if m:
            current_q = int(m.group(1))
            options_since_q[current_q] = []
            continue
        
        m = re.match(r'^[Qq](\d+)[\.:]\s+', text)
        if m:
            current_q = int(m.group(1))
            options_since_q[current_q] = []
            continue
        
        if current_q is not None:
            m_opt = re.match(r'^([A-H])[\.\)]\s', text)
            if m_opt:
                letter = m_opt.group(1)
                bold_texts = ''.join(r.text for r in para.runs if r.bold).strip()
                
                options_since_q.setdefault(current_q, []).append({
                    'letter': letter,
                    'full_text': text,
                    'bold': bold_texts,
                })
                
                # Check for bold answer
                if bold_texts and current_q not in answers:
                    m_boldopt = re.match(r'^([A-H])[\.\)]', bold_texts)
                    if m_boldopt:
                        if not re.match(r'^\d+\.\s', bold_texts):
                            answers[current_q] = m_boldopt.group(1)
                
                # Check for CORRECT marker
                if current_q not in answers:
                    if re.search(r'(?<!\bIN)CORRECT[):]', bold_texts, re.IGNORECASE):
                        answers[current_q] = letter
    
    # Also check full text for CORRECT/INCORRECT markers
    for para in doc.paragraphs:
        text = para.text.strip()
        bold_texts = ''.join(r.text for r in para.runs if r.bold).strip()
        
        m_ans = re.match(r'^(?:Answer|ANSWER)\s*:\s*([A-H])', bold_texts)
        if m_ans:
            letter = m_ans.group(1)
            for qnum in sorted(options_since_q.keys(), reverse=True):
                if qnum not in answers:
                    answers[qnum] = letter
                    break
    
    return answers

# Check specific files
for fname in ["Mini-ABT exam with answers 05 May 2017.docx",
              "Mini-ABT exam with answers for 11 August 2017.docx",
              "Mini-ABT examination with answers 02 June 2017.docx",
              "Mini-ABT exam with answers for 22 Sept. 2017 (corrected.ver.2).docx"]:
    fpath = os.path.join(EXAM_DIR, fname)
    answers = extract_answers_v4(fpath)
    
    # Show specific qnums
    target_qs = {12, 14, 15, 26, 29, 30, 16, 17, 18, 19, 20, 27, 28, 36, 37, 38, 39, 40, 11, 13, 25, 1, 32, 33}
    found = {q: a for q, a in answers.items() if q in target_qs}
    missing = [q for q in target_qs if q not in answers and q <= max(answers.keys() | {0})]
    
    print(f"\n{fname}:")
    print(f"  Found: {found}")
    print(f"  Missing from targets: {missing}")
    print(f"  Total answers: {len(answers)}")
    
    # Also show all answers for 11 August and 22 Sept
    if '11 August' in fname:
        print(f"  ALL answers: {dict(sorted(answers.items()))}")
    if '22 Sept' in fname:
        print(f"  Q32: {answers.get(32, 'NOT FOUND')}")
        print(f"  Q33: {answers.get(33, 'NOT FOUND')}")
        # Check what's around Q32-33
        print(f"  Q30-35: {{k: answers.get(k) for k in range(30, 36)}}")
