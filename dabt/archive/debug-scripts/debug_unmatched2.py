#!/usr/bin/env python3
"""Find unmatched DB questions in docx files via fuzzy text match."""
import os, re, sqlite3
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"

conn = sqlite3.connect(DB_PATH)
conn.row_factory = sqlite3.Row
cur = conn.cursor()
cur.execute("SELECT id, question_number_in_source, question_text FROM questions WHERE source_file_id=4 ORDER BY id")
db_questions = [dict(r) for r in cur.fetchall()]
conn.close()

unmatched_ids = [
    'DABT-3433', 'DABT-3434', 'DABT-3445', 'DABT-3446', 'DABT-3447',
    'DABT-3470', 'DABT-3480', 'DABT-3481', 'DABT-3482', 'DABT-3483',
    'DABT-3490', 'DABT-3491', 'DABT-3492', 'DABT-3493', 'DABT-3494',
    'DABT-3554', 'DABT-3555'
]

all_files = sorted(os.listdir(EXAM_DIR))
with_ans_files = [f for f in all_files if ('with answers' in f.lower() or 'examination with answers' in f.lower() or 'exam answers' in f.lower())]

for db_id in unmatched_ids:
    db_q = next(q for q in db_questions if q['id'] == db_id)
    db_text = db_q['question_text'].strip()
    qnum = db_q['question_number_in_source']
    
    print(f"\n{'='*60}")
    print(f"[{db_id}] Q{qnum}: {db_text[:80]}")
    
    best_match = None
    best_score = 0
    
    for fname in with_ans_files:
        fpath = os.path.join(EXAM_DIR, fname)
        doc = Document(fpath)
        
        for para in doc.paragraphs:
            ptext = para.text.strip()
            m = re.match(r'^(\d+)[\.\)]\s+(.*)', ptext)
            if not m:
                m = re.match(r'^[Qq](\d+)[\.:]\s+(.*)', ptext)
            if not m:
                continue
            
            eqnum = int(m.group(1))
            etext = m.group(2).strip()
            
            # Count common prefix
            db_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', db_text.lower())).strip()
            e_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', etext.lower())).strip()
            
            score = 0
            for i in range(min(len(db_norm), len(e_norm))):
                if db_norm[i] == e_norm[i]:
                    score += 1
                else:
                    break
            
            if score > best_score:
                best_score = score
                best_match = (fname, eqnum, etext)
    
    if best_match and best_score > 10:
        fname, eqnum, etext = best_match
        print(f"  BEST MATCH: {fname} Q{eqnum} (score={best_score})")
        print(f"  Exam text: {etext[:80]}")
        
        # Now find the answer for this exam question
        doc = Document(os.path.join(EXAM_DIR, fname))
        current_q = None
        for para in doc.paragraphs:
            text = para.text.strip()
            m = re.match(r'^(\d+)[\.\)]\s+', text)
            if m:
                current_q = int(m.group(1))
            m = re.match(r'^[Qq](\d+)[\.:]\s+', text)
            if m:
                current_q = int(m.group(1))
            
            if current_q != eqnum:
                continue
            
            bold_texts = [r.text for r in para.runs if r.bold]
            if bold_texts:
                full_bold = ''.join(bold_texts).strip()
                m_ans = re.match(r'^(?:Answer|ANSWER)\s*:\s*([A-H])', full_bold)
                if m_ans:
                    print(f"  ANSWER: {m_ans.group(1)} (ANSWER format)")
                    break
                m_opt = re.match(r'^([A-H])[\.\)]', full_bold)
                if m_opt and not re.match(r'^\d+\.\s', full_bold) and not re.match(r'^[Qq]\d+', full_bold):
                    print(f"  ANSWER: {m_opt.group(1)} (bold option format)")
                    break
    else:
        print(f"  NO GOOD MATCH (best_score={best_score})")
