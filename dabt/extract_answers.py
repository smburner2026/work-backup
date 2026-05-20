#!/usr/bin/env python3
"""
Extract correct answer letters from Chapter Tests with Answers docx files
and update the DABT database.

Usage: python3 extract_answers.py
"""

import os
import re
import sqlite3
import json
from docx import Document

# Paths
ANSWERS_DIR = "/root/dabt-curated/Chapter_Tests/Tests_with_Answers"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_PATH = "/root/work/dabt/chapter_test_answers.json"
SOURCE_FILE_ID = 3

# List of answer docx files (in order matching DB chapter order)
ANSWER_FILES = [
    "Chapter 11 - Toxic Responses of the Blood with Answers.docx",
    "Chapter 12A - Toxic Responses of the Immune System with Answers.docx",
    "Chapter 12B - Toxic Responses of the Immune System with Answers.docx",
    "Chapter 13 - Toxic Responses of the Liver with Answers.docx",
    "Chapter 14 - Toxic Responses of the Kidney with Answers.docx",
    "Chapter 15 - Toxic Responses of the Respiratory System with Answers.docx",
    "Chapter 16 - Toxic Responses of the Nervous System with Answers.docx",
    "Chapter 17 - Toxic Responses of the Ocular and Visual System with Answers.docx",
    "Chapter 18 - Toxic Responses of the Heart and Vascular System with Answers.docx",
    "Chapter 19 - Toxic Responses of the Skin with Answers.docx",
    "Chapter 2 - General Principles with Answers.docx",
    "Chapter 9 - Genetic Toxicology with Answers.docx",
    "Chapter 23 - Toxic Effects of Metals with Answers.docx",
    "Chapter 24 - Toxic Effects of Solvents and Vapors with Answers.docx",
    "Chapter 3 - Mechanisms of Toxicity with Answers.docx",
    "Chapter 30 - Food Toxicology with Answers.docx",
    "Chapter 5 - Absorption Distribution and Excretion of Toxicants with Answers.docx",
    "Chapter 6A - Biotransformation of Xenobiotics with Answers.docx",
    "Chapter 6B - Biotransformation of Xenobiotics with Answers.docx",
    "Chapter 7 - Toxicokinetics with Answers.docx",
    "Chapter 8 - Chemical Carcinogens with Answers.docx",
    "Chapters 10 and 20 - Developmental and Reproductive Tox with Answers.docx",
    "Chapters 26 and 27 - Animal Venoms and Plants with Answers.docx",
    "Chapters 28 and 29 - Air Pollution and Ecotoxicology with Answers.docx",
    "Chapter 21 - Toxic Responses of the Endocrine System with Answers.docx",
]


def extract_answers_from_docx(filepath):
    """
    Extract question_number -> answer_letter mappings from a docx file.
    Returns a dict: {question_number: answer_letter}
    """
    doc = Document(filepath)
    
    # Parse the document structure
    # Questions are paragraphs starting with "N. text"
    # Options are paragraphs starting with "A.", "B.", etc.
    # The correct answer option is bolded
    
    answers = {}  # {qnum: answer_letter}
    current_qnum = None
    options = []  # List of (option_letter, is_bold)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this is a question line (starts with number followed by period)
        q_match = re.match(r'^(\d+)\.\s', text)
        if q_match:
            # Save previous question's answer
            if current_qnum is not None and options:
                for opt_letter, is_bold in options:
                    if is_bold:
                        answers[current_qnum] = opt_letter
                        break
            
            current_qnum = int(q_match.group(1))
            options = []
            continue
        
        # Check if this is an option line (starts with letter followed by period)
        opt_match = re.match(r'^([A-Z])\.\s', text)
        if opt_match and current_qnum is not None:
            opt_letter = opt_match.group(1)
            
            # Check if this paragraph has bold runs
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            
            options.append((opt_letter, is_bold))
            
            # If exactly this option is bold, record immediately
            # (handles case where explanation lines aren't part of options)
            continue
        
        # Non-option, non-question line - could be explanation or continuation
        # Check if this is bold (some answer options span multiple runs)
        if current_qnum is not None and options:
            # Check if this bold paragraph could be an option continuation
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            if is_bold and options:
                # This might be a continuation of a bold option
                # Check last option
                last_opt = options[-1]
                if last_opt[1]:  # already bold
                    pass  # continuation, already recorded
                else:
                    # This bold text could be an option that wasn't caught
                    # Let's check if it starts with A. B. C. etc.
                    cont_match = re.match(r'^([A-Z])\.\s', text)
                    if cont_match:
                        options.append((cont_match.group(1), True))
    
    # Don't forget the last question
    if current_qnum is not None and options:
        for opt_letter, is_bold in options:
            if is_bold:
                answers[current_qnum] = opt_letter
                break
    
    return answers


def get_db_chapter_ranges(db_path, source_file_id):
    """
    Get DB chapter ranges (consecutive ID groups where question_number_in_source resets to 1).
    Returns list of (start_id, end_id, start_qtext).
    """
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    
    cur.execute("""
        SELECT id, question_number_in_source, substr(question_text, 1, 100) as qtext
        FROM questions
        WHERE source_file_id = ?
        ORDER BY id
    """, (source_file_id,))
    
    rows = cur.fetchall()
    conn.close()
    
    chapters = []
    current_chapter = None
    
    for row in rows:
        qid, qnum, qtext = row
        if qnum == 1:
            if current_chapter is not None:
                chapters.append(current_chapter)
            current_chapter = {'start_id': qid, 'end_id': qid, 'qtext': qtext, 'questions': []}
        if current_chapter is not None:
            current_chapter['end_id'] = qid
            current_chapter['questions'].append(row)
    
    if current_chapter is not None:
        chapters.append(current_chapter)
    
    return chapters


def match_chapters_to_files(chapters, answers_dir, answer_files):
    """
    Match DB chapter ranges to answer docx files by comparing first question text.
    Returns list of (chapter_index, filepath, extracts) for matching pairs.
    """
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    matched = []
    unmatched_files = list(answer_files)
    unmatched_chapters = list(range(len(chapters)))
    
    for ch_idx, ch in enumerate(chapters):
        db_qtext = ch['qtext'].strip().lower()
        best_match = None
        best_score = 0
        
        for fname in list(unmatched_files):
            filepath = os.path.join(answers_dir, fname)
            if not os.path.exists(filepath):
                continue
            
            answers = extract_answers_from_docx(filepath)
            if not answers:
                continue
            
            # Get the first question text from the file
            doc = Document(filepath)
            first_qtext = None
            for para in doc.paragraphs:
                text = para.text.strip()
                if re.match(r'^1\.\s', text):
                    first_qtext = text[3:].strip().lower()  # Remove "1. "
                    break
            
            if first_qtext:
                # Compare how well the first question text matches
                # Use simple word overlap
                db_words = set(db_qtext.split()[:10])
                file_words = set(first_qtext.split()[:10])
                overlap = len(db_words & file_words)
                ratio = overlap / max(len(db_words | file_words), 1)
                
                if ratio > best_score:
                    best_score = ratio
                    best_match = (ch_idx, filepath, answers, fname)
        
        if best_match and best_score > 0.3:
            matched.append(best_match)
            ch_idx_m, fp, answers, fname = best_match
            if fname in unmatched_files:
                unmatched_files.remove(fname)
            if ch_idx_m in unmatched_chapters:
                unmatched_chapters.remove(ch_idx_m)
    
    conn.close()
    return matched, unmatched_files, unmatched_chapters


def match_by_text_and_number(answers, ch_start_id, ch_end_id, db_path):
    """
    Match extracted answers to DB questions by comparing question text and number.
    Returns dict mapping {db_id: answer_letter}.
    """
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    
    # Get all questions in this chapter range
    cur.execute("""
        SELECT id, question_number_in_source, question_text
        FROM questions
        WHERE id BETWEEN ? AND ?
        ORDER BY id
    """, (ch_start_id, ch_end_id))
    
    db_questions = cur.fetchall()
    conn.close()
    
    updates = {}
    
    for qnum, answer_letter in answers.items():
        # Find the DB question(s) with this question_number_in_source
        matching_qs = [q for q in db_questions if q[1] == qnum]
        
        if len(matching_qs) == 1:
            # Direct match
            updates[matching_qs[0][0]] = answer_letter
        elif len(matching_qs) > 1:
            # Multiple questions with same number - need text matching
            # Usually shouldn't happen within a single chapter
            # But let's handle it by just picking the right one
            updates[matching_qs[0][0]] = answer_letter
            print(f"  WARNING: Multiple DB questions with Q{qnum} in range {ch_start_id}-{ch_end_id}")
    
    return updates


def main():
    print("=" * 60)
    print("DABT Chapter Test Answer Extraction")
    print("=" * 60)
    
    # Step 1: Get DB chapter ranges
    print("\n[Step 1] Getting DB chapter ranges...")
    chapters = get_db_chapter_ranges(DB_PATH, SOURCE_FILE_ID)
    print(f"  Found {len(chapters)} chapter ranges in DB")
    for i, ch in enumerate(chapters):
        print(f"  Ch{i}: {ch['start_id']} - {ch['end_id']} ({len(ch['questions'])} Qs) Q1: {ch['qtext'][:60]}...")
    
    # Step 2: Extract answers from each docx file
    print("\n[Step 2] Extracting answers from docx files...")
    
    # Process each answer file
    all_extracts = {}  # {filename: {qnum: answer_letter}}
    
    for fname in ANSWER_FILES:
        filepath = os.path.join(ANSWERS_DIR, fname)
        if not os.path.exists(filepath):
            print(f"  SKIP (not found): {fname}")
            continue
        
        answers = extract_answers_from_docx(filepath)
        all_extracts[fname] = answers
        qnums = sorted(answers.keys())
        print(f"  {fname}: {len(answers)} answers found (Q{qnums[0]}-Q{qnums[-1]})")
    
    # Step 3: Map chapters to files and collect updates
    print("\n[Step 3] Matching chapters to files and collecting updates...")
    
    # We need to figure out the correct ordering of docx files to DB chapters
    # Let's match by first question text
    
    # Build a mapping: for each DB chapter, find which docx file it corresponds to
    # by looking at the first question text
    
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    chapter_file_map = {}  # {chapter_index: filename}
    
    # Get the first question text for each chapter from DB
    for ch_idx, ch in enumerate(chapters):
        db_qtext = ch['qtext'].strip().lower()
        
        for fname in ANSWER_FILES:
            if fname in chapter_file_map.values():
                continue  # already matched
            
            filepath = os.path.join(ANSWERS_DIR, fname)
            if not os.path.exists(filepath):
                continue
            
            doc = Document(filepath)
            first_qtext = None
            for para in doc.paragraphs:
                text = para.text.strip()
                m = re.match(r'^1\.\s(.+)', text)
                if m:
                    first_qtext = m.group(1).strip().lower()
                    break
            
            if first_qtext:
                # Compare first 40 chars
                if db_qtext[:40] == first_qtext[:40]:
                    chapter_file_map[ch_idx] = fname
                    print(f"  Matched Ch{ch_idx} ({ch['start_id']}-{ch['end_id']}) -> {fname}")
                    break
                # Also try shorter comparison
                elif db_qtext[:30] == first_qtext[:30]:
                    chapter_file_map[ch_idx] = fname
                    print(f"  Matched Ch{ch_idx} ({ch['start_id']}-{ch['end_id']}) -> {fname} (30-char match)")
                    break
    
    print(f"\n  Matched {len(chapter_file_map)} / {len(chapters)} chapters to files")
    unmatched_chs = [i for i in range(len(chapters)) if i not in chapter_file_map]
    if unmatched_chs:
        print(f"  UNMATCHED chapters: {unmatched_chs}")
    
    # Step 4: Build complete update mapping
    all_updates = {}  # {question_id: answer_letter}
    
    for ch_idx, fname in chapter_file_map.items():
        ch = chapters[ch_idx]
        answers = all_extracts.get(fname, {})
        
        if not answers:
            print(f"  WARNING: No answers extracted for {fname}")
            continue
        
        # Get all DB questions in this chapter range
        cur.execute("""
            SELECT id, question_number_in_source, question_text
            FROM questions
            WHERE id BETWEEN ? AND ?
            ORDER BY id
        """, (ch['start_id'], ch['end_id']))
        
        db_questions = cur.fetchall()
        
        # Build a lookup by question number
        db_by_qnum = {}
        for q in db_questions:
            qid, qnum, qtext = q
            if qnum not in db_by_qnum:
                db_by_qnum[qnum] = []
            db_by_qnum[qnum].append((qid, qtext))
        
        # Match each extracted answer to a DB question
        for qnum, answer_letter in answers.items():
            if qnum in db_by_qnum:
                matches = db_by_qnum[qnum]
                if len(matches) == 1:
                    all_updates[matches[0][0]] = answer_letter
                else:
                    # Multiple matches - try to match by exact question text
                    # Get the question text from the docx for this qnum
                    filepath = os.path.join(ANSWERS_DIR, fname)
                    doc = Document(filepath)
                    docx_qtext = None
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        m = re.match(rf'^{qnum}\.\s(.+)', text)
                        if m:
                            docx_qtext = m.group(1).strip().lower()
                            break
                    
                    if docx_qtext:
                        # Find the best match
                        best_match = None
                        best_score = 0
                        for qid, qtext in matches:
                            # Compare first 60 chars
                            qtext_lower = qtext.strip().lower()
                            score = sum(1 for a, b in zip(docx_qtext, qtext_lower) if a == b)
                            if score > best_score:
                                best_score = score
                                best_match = qid
                        if best_match:
                            all_updates[best_match] = answer_letter
                        else:
                            print(f"  WARNING: Cannot match Q{qnum} in {fname}")
                    else:
                        # Just use first match
                        all_updates[matches[0][0]] = answer_letter
            else:
                print(f"  WARNING: Q{qnum} from {fname} not found in DB chapter range")
    
    conn.close()
    
    print(f"\n[Step 4] Collected {len(all_updates)} updates")
    
    # Step 5: Save mapping to JSON
    mapping = {
        'source_file_id': SOURCE_FILE_ID,
        'updates': {k: v for k, v in sorted(all_updates.items())},
        'total_updates': len(all_updates)
    }
    
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    with open(OUTPUT_PATH, 'w') as f:
        json.dump(mapping, f, indent=2)
    print(f"  Saved mapping to {OUTPUT_PATH}")
    
    # Step 6: Update database
    print("\n[Step 5] Updating database...")
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # First check which already have answers
    cur.execute("""
        SELECT id, correct_answer_letter FROM questions 
        WHERE source_file_id = ? AND id IN ({})
    """.format(','.join('?' for _ in all_updates.keys())), 
        (SOURCE_FILE_ID, *list(all_updates.keys())))
    
    existing = {row[0]: row[1] for row in cur.fetchall()}
    
    update_count = 0
    skip_count = 0
    conflict_count = 0
    
    for qid, answer_letter in sorted(all_updates.items()):
        if qid in existing:
            existing_letter = existing[qid]
            if existing_letter and existing_letter != answer_letter:
                print(f"  CONFLICT: {qid} has existing answer '{existing_letter}', extracted '{answer_letter}'")
                conflict_count += 1
                continue
            elif existing_letter == answer_letter:
                skip_count += 1
                continue
        
        cur.execute(
            "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
            (answer_letter, qid)
        )
        if cur.rowcount > 0:
            update_count += 1
    
    conn.commit()
    conn.close()
    
    print(f"\n{'=' * 60}")
    print(f"SUMMARY")
    print(f"{'=' * 60}")
    print(f"  Total updates in mapping: {len(all_updates)}")
    print(f"  Actually updated in DB: {update_count}")
    print(f"  Already had same answer: {skip_count}")
    print(f"  Conflicts (existing != extracted): {conflict_count}")
    print(f"  Mapping saved to: {OUTPUT_PATH}")
    print(f"{'=' * 60}")


if __name__ == '__main__':
    main()
