#!/usr/bin/env python3
"""
Extract correct answer letters from Kristen Topic Tests "with answers" docx files
and update the DABT database.
"""

import json
import re
import sqlite3
import sys
from docx import Document

DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
SOURCE_DIR = "/root/dabt-curated/Practice_Tests_by_Topic/Kristen_Topic_Tests"

# Mapping of source files to DB topic and question ranges
# Each file has its own internal numbering starting from 1
FILES_MAP = [
    {
        "filename": "Ecotox exam with answers 28 July 2017.docx",
        "topic": "General Toxicology",
        "db_ids": None,  # Will be populated from DB
        "q_start": 1,
        "q_end": 25,
    },
    {
        "filename": "Forensic-analytical-clinical exam with answers 28 July 2017.docx",
        "topic": "Drugs & Therapeutics – Toxicology",
        "db_ids": None,
        "q_start": 1,
        "q_end": 53,
    },
    {
        "filename": "Pesticides examination with answers 08 June 2017.docx",
        "topic": "Pesticides – Insecticides",
        "db_ids": None,
        "q_start": 1,
        "q_end": 50,
    },
    {
        "filename": "Risk Assessement Overview examination with answers 05 August 2017.docx",
        "topic": "Risk Assessment & Regulatory",
        "db_ids": None,
        "q_start": 1,
        "q_end": 19,
    },
]

def extract_answers_from_docx(filepath):
    """
    Extract question_number -> answer_letter from a "with answers" docx file.
    
    The pattern: The question paragraph contains the question text and ALL answer options.
    The correct answer option has bold text (the answer letter like "A." is bold).
    
    Returns: dict of {question_number: answer_letter}
    """
    doc = Document(filepath)
    answers = {}
    
    current_qnum = None
    answer_letter = None
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        # Check if this paragraph starts a new question (starts with digit)
        # Pattern: "1. ..." or "10. ..." etc.
        m = re.match(r'^(\d+)\.\s', text)
        if m:
            # Save previous question's answer if we have one
            if current_qnum is not None and answer_letter is not None:
                answers[current_qnum] = answer_letter
            
            current_qnum = int(m.group(1))
            answer_letter = None
            
            # Check for bold within this paragraph to find correct answer
            for r in p.runs:
                if r.bold and r.text.strip():
                    # Look for answer letter pattern like "A." or "B." etc.
                    bold_text = r.text.strip()
                    am = re.match(r'^([A-H])\.\s*$', bold_text)
                    if am:
                        answer_letter = am.group(1)
                        break
                    # Also check if the bold text starts with letter.
                    # In some cases, the bold includes the option text too.
                    am2 = re.match(r'^([A-H])\.?\s', bold_text)
                    if am2:
                        answer_letter = am2.group(1)
                        break
    
    # Don't forget the last question
    if current_qnum is not None and answer_letter is not None:
        answers[current_qnum] = answer_letter
    
    # ALSO check for standalone bold answer lines (some answers might be in separate paragraphs)
    # This happens when the answer option is on its own line
    current_qnum_2 = None
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Find question numbers
        m = re.match(r'^(\d+)\.\s', text)
        if m:
            current_qnum_2 = int(m.group(1))
            # Check for bold in the question paragraph
            for r in p.runs:
                if r.bold and r.text.strip():
                    bold_text = r.text.strip()
                    am = re.match(r'^([A-H])\.?\s*', bold_text)
                    if am:
                        answers[current_qnum_2] = am.group(1)
            continue
        
        # Check for paragraphs that are just an answer option (e.g., "A. text")
        # with the letter bold
        if current_qnum_2 is not None:
            am = re.match(r'^([A-H])\.\s', text)
            if am:
                letter = am.group(1)
                for r in p.runs:
                    if r.bold and r.text.strip():
                        bold_text = r.text.strip()
                        bm = re.match(r'^([A-H])\.?\s*', bold_text)
                        if bm and bm.group(1) == letter:
                            answers[current_qnum_2] = letter
                            break
    
    return answers


def get_db_questions(topic):
    """Get questions from DB for a given topic/range."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT q.id, q.question_number_in_source, q.correct_answer_letter, qt.topic
        FROM questions q
        JOIN question_topics qt ON q.id = qt.question_id
        WHERE q.source_file_id = 5 AND qt.topic = ?
        ORDER BY q.question_number_in_source
    """, (topic,))
    
    rows = cursor.fetchall()
    conn.close()
    return rows


def main():
    all_answers = {}  # file -> {qnum: letter}
    
    for file_info in FILES_MAP:
        filepath = f"{SOURCE_DIR}/{file_info['filename']}"
        print(f"\n{'='*60}")
        print(f"Processing: {file_info['filename']}")
        print(f"  Topic: {file_info['topic']}")
        print(f"  Range: Q{file_info['q_start']}-Q{file_info['q_end']}")
        
        answers = extract_answers_from_docx(filepath)
        all_answers[file_info['filename']] = answers
        print(f"  Extracted {len(answers)} answers")
        
        # Show what we found
        for qnum in sorted(answers.keys()):
            print(f"    Q{qnum}: {answers[qnum]}")
    
    # Save raw extraction
    with open("/root/work/dabt/kristen_topic_answers_raw.json", "w") as f:
        json.dump(all_answers, f, indent=2)
    
    print(f"\n{'='*60}")
    print("Raw extraction saved to /root/work/dabt/kristen_topic_answers_raw.json")


if __name__ == "__main__":
    main()
