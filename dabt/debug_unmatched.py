#!/usr/bin/env python3
"""Debug unmatched questions - find which exam they belong to."""
import os, re, sqlite3
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"

# Get DB questions
conn = sqlite3.connect(DB_PATH)
conn.row_factory = sqlite3.Row
cur = conn.cursor()
cur.execute("SELECT id, question_number_in_source, question_text FROM questions WHERE source_file_id=4 ORDER BY id")
db_questions = [dict(r) for r in cur.fetchall()]
conn.close()

# The unmatched IDs from the output
unmatched_ids = [
    'DABT-3433', 'DABT-3434', 'DABT-3445', 'DABT-3446', 'DABT-3447',
    'DABT-3470', 'DABT-3480', 'DABT-3481', 'DABT-3482', 'DABT-3483',
    'DABT-3490', 'DABT-3491', 'DABT-3492', 'DABT-3493', 'DABT-3494',
    'DABT-3554', 'DABT-3555'
]

# For each unmatched question, search for its text in all docx files
all_files = sorted(os.listdir(EXAM_DIR))
docx_files = [f for f in all_files if f.endswith('.docx')]

for db_id in unmatched_ids:
    db_q = next(q for q in db_questions if q['id'] == db_id)
    db_text = db_q['question_text'].strip()
    qnum = db_q['question_number_in_source']
    
    print(f"\n{'='*60}")
    print(f"[{db_id}] Q{qnum}: {db_text[:80]}")
    
    for fname in docx_files:
        fpath = os.path.join(EXAM_DIR, fname)
        doc = Document(fpath)
        
        for para in doc.paragraphs:
            ptext = para.text.strip()
            # Check if this paragraph contains the DB question text as a substring
            if len(ptext) > 20 and (db_text[:30].lower() in ptext.lower() or ptext.lower().startswith(db_text[:20].lower())):
                # Check if it's a question line (starts with number)
                if re.match(r'^\d+[\.\)]\s', ptext) or re.match(r'^[Qq]\d+[\.:]', ptext):
                    # Find the answer for this question
                    print(f"  Found in {fname}")
                    print(f"  Text: {ptext[:100]}")
                    
                    # Check bold runs
                    bold_runs = [r.text for r in para.runs if r.bold]
                    if bold_runs:
                        bold_text = ''.join(bold_runs).strip()
                        print(f"  Bold: '{bold_text[:80]}'")
                    
                    # Find the answer option in the same file
                    ans = find_answer_in_file(doc, ptext)
                    if ans:
                        print(f"  Answer: {ans}")
                    break
    
    # Also try searching more broadly
    found = False
    for fname in docx_files:
        if 'with answers' not in fname.lower() and 'examination with answers' not in fname.lower() and 'exam answers' not in fname.lower():
            continue
        fpath = os.path.join(EXAM_DIR, fname)
        doc = Document(fpath)
        for para in doc.paragraphs:
            ptext = para.text.strip()
            # More flexible matching
            if len(ptext) > 30:
                # Check if significant portion matches
                words_db = set(re.findall(r'\w+', db_text.lower()))
                words_pt = set(re.findall(r'\w+', ptext.lower()))
                common = len(words_db & words_pt)
                if common >= 5 and common / max(len(words_db), 1) > 0.5:
                    if re.match(r'^\d+[\.\)]\s', ptext) or re.match(r'^[Qq]\d+[\.:]', ptext):
                        print(f"  Also found (fuzzy) in {fname}: {ptext[:100]}")
                        found = True
                        break
        if found:
            break

def find_answer_in_file(doc, question_text):
    """Find the answer letter for a question in a docx file."""
    # Extract question number
    m = re.match(r'^(\d+)', question_text)
    m2 = re.match(r'^[Qq](\d+)', question_text)
    qnum = None
    if m:
        qnum = int(m.group(1))
    elif m2:
        qnum = int(m2.group(1))
    
    if qnum is None:
        return None
    
    # Look for bolded answer options
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        bold_runs = [r.text for r in para.runs if r.bold]
        if not bold_runs:
            continue
        
        full_bold = ''.join(bold_runs).strip()
        
        # Check for "ANSWER: X"
        ans_m = re.match(r'^(?:Answer|ANSWER)\s*:\s*([A-H])', full_bold)
        if ans_m:
            # Need to associate with correct question
            pass
        
        # Check for option format
        opt_m = re.match(r'^([A-H])[\.\)]\s', full_bold)
        if opt_m and not re.match(r'^\d+\.\s', full_bold) and not re.match(r'^[Qq]\d+', full_bold):
            # This is a bolded option - but which question does it belong to?
            pass
