#!/usr/bin/env python3
"""
v5 extractor: handles CORRECT/INCORRECT markers on paragraphs FOLLOWING the option.
Also handles all other formats.
"""

import os, re, json, sqlite3
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_JSON = "/root/work/dabt/kristen_mini_answers.json"

def extract_answers_v5(doc_path):
    doc = Document(doc_path)
    answers = {}
    
    # Collect all paragraphs with their properties
    paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        bold = ''.join(r.text for r in para.runs if r.bold).strip()
        paras.append({'text': text, 'bold': bold, 'runs': para.runs})
    
    # First pass: identify question numbers
    # Track the last seen question for each option
    current_q = None
    last_option_letter = None  # Track the last seen option letter
    
    i = 0
    while i < len(paras):
        p = paras[i]
        text = p['text']
        bold = p['bold']
        
        # Question number detection
        m = re.match(r'^(\d+)[\.\)]\s+', text)
        if m:
            current_q = int(m.group(1))
            last_option_letter = None
            i += 1
            continue
        
        m = re.match(r'^[Qq](\d+)[\.:]\s+', text)
        if m:
            current_q = int(m.group(1))
            last_option_letter = None
            i += 1
            continue
        
        # Option letter detection
        m_opt = re.match(r'^([A-H])[\.\)]', text)
        if m_opt and current_q is not None:
            # This is an option line
            letter = m_opt.group(1)
            last_option_letter = letter
            
            # Check if THIS paragraph has bold with answer indicator
            if current_q not in answers:
                # Check for CORRECT in bold (not INCORRECT)
                if re.search(r'(?<!\bIN)CORRECT[):]', bold, re.IGNORECASE):
                    answers[current_q] = letter
                    i += 1
                    continue
                
                # Check for bold option letter
                m_boldopt = re.match(r'^([A-H])[\.\)]', bold)
                if m_boldopt and not re.match(r'^\d+\.\s', bold) and not re.match(r'^[Qq]\d+', bold):
                    answers[current_q] = letter
                    i += 1
                    continue
            
            # Check NEXT paragraph for answer indicator
            if i + 1 < len(paras):
                next_bold = paras[i+1]['bold']
                next_text = paras[i+1]['text']
                
                # Skip if next para is a new question
                if not re.match(r'^\d+[\.\)]\s', next_text) and not re.match(r'^[Qq]\d+', next_text):
                    if current_q not in answers:
                        if re.search(r'(?<!\bIN)CORRECT[):]', next_bold, re.IGNORECASE):
                            answers[current_q] = letter
                        elif re.search(r'\bCORRECT\b', next_bold, re.IGNORECASE) and not re.search(r'\bINCORRECT\b', next_bold, re.IGNORECASE):
                            answers[current_q] = letter
            
            i += 1
            continue
        
        # Check for "ANSWER: X" format (bold)
        m_ans = re.match(r'^(?:Answer|ANSWER)\s*:\s*([A-H])', bold)
        if m_ans and current_q is not None and current_q not in answers:
            answers[current_q] = m_ans.group(1)
            i += 1
            continue
        
        i += 1
    
    return answers

def extract_question_texts(doc_path):
    doc = Document(doc_path)
    qtexts = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m = re.match(r'^(\d+)[\.\)]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            qtexts[qnum] = m.group(2).strip()
        m = re.match(r'^[Qq](\d+)[\.:]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            if qnum not in qtexts:
                qtexts[qnum] = m.group(2).strip()
    return qtexts

def normalize(t):
    t = re.sub(r'\s+', ' ', t).strip().lower()
    t = re.sub(r'[^\w\s]', '', t)
    return t

# Process all files
all_files = sorted(os.listdir(EXAM_DIR))
answer_files = [f for f in all_files if ('with answers' in f.lower() or 'examination with answers' in f.lower() or 'exam answers' in f.lower())]

exam_data = {}
for fname in answer_files:
    fpath = os.path.join(EXAM_DIR, fname)
    answers_v5 = extract_answers_v5(fpath)
    qtexts = extract_question_texts(fpath)
    exam_data[fname] = {'answers': answers_v5, 'qtexts': qtexts}
    
    missing = [q for q in sorted(qtexts.keys()) if q not in answers_v5]
    print(f"{fname}: {len(answers_v5)} answers, {len(qtexts)} texts" + 
          (f"  MISSING: {missing}" if missing else ""))

# Get DB questions
conn = sqlite3.connect(DB_PATH)
conn.row_factory = sqlite3.Row
cur = conn.cursor()
cur.execute("SELECT id, question_number_in_source, question_text FROM questions WHERE source_file_id=4 ORDER BY id")
db_questions = [dict(r) for r in cur.fetchall()]
conn.close()

# For each DB question, find best matching exam question by text
id_to_answer = {}
unmatched = []

for db_q in db_questions:
    db_text = db_q['question_text']
    db_id = db_q['id']
    db_norm = normalize(db_text)
    
    best_match = None
    best_score = 0
    
    for fname, data in exam_data.items():
        answers = data['answers']
        qtexts = data['qtexts']
        
        for qnum, answer_letter in answers.items():
            if qnum not in qtexts:
                continue
            
            exam_text = qtexts[qnum]
            exam_norm = normalize(exam_text)
            
            score = 0
            for i in range(min(len(db_norm), len(exam_norm))):
                if db_norm[i] == exam_norm[i]:
                    score += 1
                else:
                    break
            
            if score > best_score:
                best_score = score
                best_match = (fname, qnum, answer_letter)
    
    if best_match and best_score >= 5:
        fname, qnum, answer = best_match
        id_to_answer[db_id] = {
            'answer_letter': answer,
            'exam': fname,
            'qnum': qnum,
            'match_score': best_score
        }
    else:
        unmatched.append((db_id, db_text, best_score))

print(f"\nMatched: {len(id_to_answer)}/{len(db_questions)}")
print(f"Unmatched: {len(unmatched)}")
for db_id, text, score in unmatched:
    print(f"  [{db_id}]: {text[:60]} (score={score})")

# Save and update
with open(OUTPUT_JSON, 'w') as f:
    json.dump(id_to_answer, f, indent=2)

conn = sqlite3.connect(DB_PATH)
cursor = conn.cursor()
updated = 0
for db_id, info in id_to_answer.items():
    cursor.execute("UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
                   (info['answer_letter'], db_id))
    if cursor.rowcount > 0:
        updated += 1
conn.commit()

cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4 AND correct_answer_letter IS NOT NULL AND correct_answer_letter != ''")
filled = cursor.fetchone()[0]
print(f"\nDB Updated: {updated} ok, Questions with answers: {filled}/{len(db_questions)}")
conn.close()
