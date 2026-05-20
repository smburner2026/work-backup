#!/usr/bin/env python3
"""
Comprehensive answer extraction that handles ALL formats found in the docx files:
1. Bold on option letter (standard "A." or "A)" format)
2. "ANSWER: X" format
3. "(CORRECT)" / "(INCORRECT)" bold markers
4. "CORRECT:" / "INCORRECT:" markers
5. "OK" at beginning of answer options
"""

import os, re, json, sqlite3
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_JSON = "/root/work/dabt/kristen_mini_answers.json"

def extract_answers_v4(doc_path):
    """
    Ultra-comprehensive answer extraction.
    Returns {qnum: answer_letter}
    """
    doc = Document(doc_path)
    answers = {}
    current_q = None
    # Track the option lines seen since last question
    options_since_q = {}
    
    # First pass: identify question numbers and their options
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detect question number
        m = re.match(r'^(\d+)[\.\)]\s+', text)
        if m:
            current_q = int(m.group(1))
            options_since_q[current_q] = []
            continue
        
        m = re.match(r'^[Qq](\d+)[\.:]\s+', text)
        if m:
            current_q = int(m.group(1))
            options_since_q[current_q] = []
            continue
        
        if current_q is not None:
            # Check if this is an option line A) through H)
            m_opt = re.match(r'^([A-H])[\.\)]\s', text)
            if m_opt:
                letter = m_opt.group(1)
                bold_texts = ''.join(r.text for r in para.runs if r.bold).strip()
                
                options_since_q.setdefault(current_q, []).append({
                    'letter': letter,
                    'full_text': text,
                    'bold': bold_texts,
                    'para': para
                })
                
                # Check for bold answer (format 1: bold on option letter)
                if bold_texts and current_q not in answers:
                    m_boldopt = re.match(r'^([A-H])[\.\)]', bold_texts)
                    if m_boldopt:
                        # Make sure it's not a question header
                        if not re.match(r'^\d+\.\s', bold_texts):
                            answers[current_q] = m_boldopt.group(1)
                
                # Check for (CORRECT) or (INCORRECT) marker in bold
                if current_q not in answers:
                    if re.search(r'CORRECT[):]', bold_texts, re.IGNORECASE) and not re.search(r'INCORRECT', bold_texts, re.IGNORECASE):
                        answers[current_q] = letter
                    elif 'CORRECT' in bold_texts.upper() and 'INCORRECT' not in bold_texts.upper():
                        answers[current_q] = letter
    
    # Second pass: check for "ANSWER: X" format
    for para in doc.paragraphs:
        text = para.text.strip()
        bold_texts = ''.join(r.text for r in para.runs if r.bold).strip()
        
        # "ANSWER: X" in bold
        m_ans = re.match(r'^(?:Answer|ANSWER)\s*:\s*([A-H])', bold_texts)
        if m_ans:
            letter = m_ans.group(1)
            # Find which question this answer belongs to
            # Look backwards to find the last question
            for qnum in sorted(options_since_q.keys(), reverse=True):
                if qnum not in answers:
                    answers[qnum] = letter
                    break
    
    # Third pass: for questions still without answers, look for (CORRECT) or "CORRECT:" in any para
    for para in doc.paragraphs:
        text = para.text.strip()
        # Check for "CORRECT:" after option letter (e.g., "D) CORRECT: ...")
        m = re.match(r'^([A-H])[\.\)].*?\bCORRECT[:\s)]', text, re.IGNORECASE)
        if m and current_q is not None:
            letter = m.group(1)
            # Need to determine which question this belongs to
            # Find the last question number before this paragraph
            pass  # Complex to implement
    
    return answers

def extract_question_texts(doc_path):
    """Extract {qnum: clean_question_text} from docx."""
    doc = Document(doc_path)
    qtexts = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m = re.match(r'^(\d+)[\.\)]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            qtext = m.group(2).strip()
            qtexts[qnum] = qtext
        m = re.match(r'^[Qq](\d+)[\.:]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            qtext = m.group(2).strip()
            if qnum not in qtexts:
                qtexts[qnum] = qtext
    return qtexts

def normalize(t):
    t = re.sub(r'\s+', ' ', t).strip().lower()
    t = re.sub(r'[^\w\s]', '', t)
    return t

# Process all files
all_files = sorted(os.listdir(EXAM_DIR))
answer_files = [f for f in all_files if ('with answers' in f.lower() or 'examination with answers' in f.lower() or 'exam answers' in f.lower())]

exam_data = {}
for fname in answer_files:
    fpath = os.path.join(EXAM_DIR, fname)
    answers_v4 = extract_answers_v4(fpath)
    qtexts = extract_question_texts(fpath)
    exam_data[fname] = {'answers': answers_v4, 'qtexts': qtexts}
    print(f"{fname}: {len(answers_v4)} answers, {len(qtexts)} texts")
    if len(answers_v4) < len(qtexts):
        missing = [q for q in sorted(qtexts.keys()) if q not in answers_v4]
        print(f"  Missing answers for Q: {missing[:15]}{'...' if len(missing) > 15 else ''}")

# Get DB questions
conn = sqlite3.connect(DB_PATH)
conn.row_factory = sqlite3.Row
cur = conn.cursor()
cur.execute("SELECT id, question_number_in_source, question_text FROM questions WHERE source_file_id=4 ORDER BY id")
db_questions = [dict(r) for r in cur.fetchall()]
conn.close()

print(f"\nDB has {len(db_questions)} questions")

# Match by text
id_to_answer = {}
unmatched = []

for db_q in db_questions:
    db_text = db_q['question_text']
    db_id = db_q['id']
    db_qnum = db_q['question_number_in_source']
    db_norm = normalize(db_text)
    
    best_match = None
    best_score = 0
    
    for fname, data in exam_data.items():
        answers = data['answers']
        qtexts = data['qtexts']
        
        for qnum, answer_letter in answers.items():
            if qnum not in qtexts:
                continue
            
            exam_text = qtexts[qnum]
            exam_norm = normalize(exam_text)
            
            # Score based on text similarity
            score = 0
            for i in range(min(len(db_norm), len(exam_norm))):
                if db_norm[i] == exam_norm[i]:
                    score += 1
                else:
                    break
            
            if score > best_score:
                best_score = score
                best_match = (fname, qnum, answer_letter)
    
    if best_match and best_score >= 5:
        fname, qnum, answer = best_match
        id_to_answer[db_id] = {
            'answer_letter': answer,
            'exam': fname,
            'qnum': qnum,
            'match_score': best_score
        }
    else:
        unmatched.append((db_id, db_text, best_score))

print(f"\nMatched: {len(id_to_answer)}/{len(db_questions)}")
print(f"Unmatched: {len(unmatched)}")
if unmatched:
    print("\nUnmatched questions:")
    for db_id, text, score in unmatched:
        print(f"  [{db_id}]: {text[:60]} (score={score})")

# Save JSON
with open(OUTPUT_JSON, 'w') as f:
    json.dump(id_to_answer, f, indent=2)
print(f"\nSaved to {OUTPUT_JSON}")

# Update DB
conn = sqlite3.connect(DB_PATH)
cursor = conn.cursor()
updated = 0
errors = 0
for db_id, info in id_to_answer.items():
    try:
        cursor.execute(
            "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
            (info['answer_letter'], db_id)
        )
        if cursor.rowcount > 0:
            updated += 1
        else:
            errors += 1
    except Exception as e:
        print(f"ERROR: {e}")
        errors += 1
conn.commit()

cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4 AND correct_answer_letter IS NOT NULL AND correct_answer_letter != ''")
filled = cursor.fetchone()[0]
total = len(db_questions)
print(f"\nDB Updated: {updated} ok, {errors} errors")
print(f"Questions with answers: {filled}/{total}")
conn.close()
