#!/usr/bin/env python3
"""
Extract correct answer letters from Chapter Tests with Answers docx files
and update the DABT database.

Version 2: Override existing answers (docx is authoritative).
"""

import os
import re
import sqlite3
import json
from docx import Document

# Paths
ANSWERS_DIR = "/root/dabt-curated/Chapter_Tests/Tests_with_Answers"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_PATH = "/root/work/dabt/chapter_test_answers.json"
SOURCE_FILE_ID = 3


def extract_answers_from_docx(filepath):
    """
    Extract question_number -> answer_letter mappings from a docx file.
    Returns a dict: {question_number: answer_letter}
    """
    doc = Document(filepath)
    
    answers = {}
    current_qnum = None
    options = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this starts a new question (number followed by period)
        q_match = re.match(r'^(\d+)\.\s', text)
        if q_match:
            rest = text[q_match.end():]
            # A new question must have substantial text after the number.
            # Also, if we're already tracking a question and haven't seen any
            # options yet, short numbered items are sub-parts, not new questions.
            is_new_question = len(rest) >= 40
            if not is_new_question and current_qnum is not None and not options:
                # This is a sub-part of the current question (e.g., "1. Hydrolysis")
                pass  # Skip - still part of current question
            elif is_new_question:
                # Save previous question's answer
                if current_qnum is not None and options:
                    for opt_letter, is_bold in options:
                        if is_bold:
                            answers[current_qnum] = opt_letter
                            break
                
                current_qnum = int(q_match.group(1))
                options = []
                continue
        
        # Check if this is a multiple choice option line
        opt_match = re.match(r'^([A-Z])\.\s', text)
        if opt_match and current_qnum is not None:
            opt_letter = opt_match.group(1)
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            options.append((opt_letter, is_bold))
            continue
    
    # Save last question's answer
    if current_qnum is not None and options:
        for opt_letter, is_bold in options:
            if is_bold:
                answers[current_qnum] = opt_letter
                break
    
    return answers


def get_db_chapter_ranges(db_path, source_file_id):
    """
    Get DB chapter ranges (consecutive ID groups where question_number_in_source resets to 1).
    Returns list of dicts with start_id, end_id, qtext, questions list.
    """
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    
    cur.execute("""
        SELECT id, question_number_in_source, substr(question_text, 1, 100) as qtext
        FROM questions
        WHERE source_file_id = ?
        ORDER BY id
    """, (source_file_id,))
    
    rows = cur.fetchall()
    conn.close()
    
    chapters = []
    current_chapter = None
    
    for row in rows:
        qid, qnum, qtext = row
        if qnum == 1:
            if current_chapter is not None:
                chapters.append(current_chapter)
            current_chapter = {'start_id': qid, 'end_id': qid, 'qtext': qtext, 'questions': []}
        if current_chapter is not None:
            current_chapter['end_id'] = qid
            current_chapter['questions'].append(row)
    
    if current_chapter is not None:
        chapters.append(current_chapter)
    
    return chapters


def main():
    print("=" * 60)
    print("DABT Chapter Test Answer Extraction v2")
    print("=" * 60)
    
    # Step 1: Get DB chapter ranges
    print("\n[Step 1] Getting DB chapter ranges...")
    chapters = get_db_chapter_ranges(DB_PATH, SOURCE_FILE_ID)
    print(f"  Found {len(chapters)} chapter ranges in DB")
    
    # Step 2: Extract answers from all docx files
    print("\n[Step 2] Extracting answers from all docx files...")
    all_extracts = {}
    
    answer_files = sorted(os.listdir(ANSWERS_DIR))
    for fname in answer_files:
        if not fname.endswith('.docx') or 'with Answers' not in fname:
            continue
        filepath = os.path.join(ANSWERS_DIR, fname)
        answers = extract_answers_from_docx(filepath)
        all_extracts[fname] = answers
        qnums = sorted(answers.keys())
        if qnums:
            print(f"  {fname}: {len(answers)} answers (Q{qnums[0]}-Q{qnums[-1]})")
        else:
            print(f"  {fname}: NO answers extracted!")
    
    # Step 3: Match chapters to files by comparing first question text
    print("\n[Step 3] Matching DB chapters to docx files...")
    
    chapter_file_map = {}
    
    for ch_idx, ch in enumerate(chapters):
        db_qtext = ch['qtext'].strip().lower()
        
        for fname, answers in all_extracts.items():
            if fname in chapter_file_map.values():
                continue
            
            if not answers:
                continue
            
            filepath = os.path.join(ANSWERS_DIR, fname)
            doc = Document(filepath)
            first_qtext = None
            for para in doc.paragraphs:
                text = para.text.strip()
                m = re.match(r'^1\.\s(.+)', text)
                if m:
                    first_qtext = m.group(1).strip().lower()
                    break
            
            if first_qtext:
                if db_qtext[:40] == first_qtext[:40]:
                    chapter_file_map[ch_idx] = fname
                    print(f"  Ch{ch_idx} ({ch['start_id']}-{ch['end_id']}) <-> {fname}")
                    break
                elif db_qtext[:30] == first_qtext[:30]:
                    chapter_file_map[ch_idx] = fname
                    print(f"  Ch{ch_idx} ({ch['start_id']}-{ch['end_id']}) <-> {fname} (30-char match)")
                    break
    
    print(f"  Matched {len(chapter_file_map)} / {len(chapters)} chapters")
    unmatched = [i for i in range(len(chapters)) if i not in chapter_file_map]
    if unmatched:
        print(f"  UNMATCHED: {unmatched}")
    
    # Step 4: Build the update mapping
    print("\n[Step 4] Building update mapping...")
    
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    all_updates = {}
    match_stats = {'matched': 0, 'skipped_qnum_not_in_db': 0, 'multiple_matches': 0}
    
    for ch_idx, fname in chapter_file_map.items():
        ch = chapters[ch_idx]
        answers = all_extracts.get(fname, {})
        
        if not answers:
            continue
        
        # Get all DB questions in this chapter range
        cur.execute("""
            SELECT id, question_number_in_source, question_text
            FROM questions
            WHERE id BETWEEN ? AND ?
            ORDER BY id
        """, (ch['start_id'], ch['end_id']))
        
        db_questions = cur.fetchall()
        
        # Build lookup by qnum
        db_by_qnum = {}
        for q in db_questions:
            qid, qnum, qtext = q
            db_by_qnum.setdefault(qnum, []).append((qid, qtext))
        
        # Match each extracted answer to DB questions
        for qnum, answer_letter in answers.items():
            if qnum not in db_by_qnum:
                match_stats['skipped_qnum_not_in_db'] += 1
                continue
            
            matches = db_by_qnum[qnum]
            if len(matches) == 1:
                all_updates[matches[0][0]] = answer_letter
                match_stats['matched'] += 1
            else:
                # Multiple DB questions with same qnum in same chapter
                # Try text matching
                doc = Document(os.path.join(ANSWERS_DIR, fname))
                docx_qtext = None
                for para in doc.paragraphs:
                    text = para.text.strip()
                    m = re.match(rf'^{qnum}\.\s(.+)', text)
                    if m:
                        docx_qtext = m.group(1).strip().lower()
                        break
                
                if docx_qtext:
                    best_match = None
                    best_score = 0
                    for qid, qtext in matches:
                        qtext_lower = qtext.strip().lower()
                        # Count matching prefix characters
                        score = sum(1 for a, b in zip(docx_qtext, qtext_lower) if a == b)
                        if score > best_score:
                            best_score = score
                            best_match = qid
                    
                    if best_match:
                        all_updates[best_match] = answer_letter
                        match_stats['matched'] += 1
                        # The other match(es) will be handled separately
                        for qid, _ in matches:
                            if qid != best_match:
                                all_updates[qid] = answer_letter  # Same answer for matching sub-questions
                                match_stats['matched'] += 1
                else:
                    # Can't match by text, assign to first
                    all_updates[matches[0][0]] = answer_letter
                    match_stats['matched'] += 1
                    match_stats['multiple_matches'] += 1
    
    conn.close()
    
    print(f"  Matched: {match_stats['matched']}")
    print(f"  Skipped (qnum not in DB): {match_stats['skipped_qnum_not_in_db']}")
    print(f"  Multiple matches (text-matched): {match_stats['multiple_matches']}")
    print(f"  Total updates: {len(all_updates)}")
    
    # Step 5: Save mapping to JSON
    print("\n[Step 5] Saving mapping to JSON...")
    mapping = {
        'source_file_id': SOURCE_FILE_ID,
        'updates': {k: v for k, v in sorted(all_updates.items())},
        'total_updates': len(all_updates)
    }
    
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    with open(OUTPUT_PATH, 'w') as f:
        json.dump(mapping, f, indent=2)
    print(f"  Saved to {OUTPUT_PATH}")
    
    # Step 6: Update database - OVERRIDE existing answers
    print("\n[Step 6] Updating database (overriding existing answers)...")
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # Check current state
    cur.execute("""
        SELECT id, correct_answer_letter FROM questions 
        WHERE source_file_id = ? AND id IN ({})
    """.format(','.join('?' for _ in all_updates.keys())), 
        (SOURCE_FILE_ID, *list(all_updates.keys())))
    
    existing = {row[0]: row[1] for row in cur.fetchall()}
    
    update_count = 0
    new_count = 0
    conflict_count = 0
    same_count = 0
    
    for qid, answer_letter in sorted(all_updates.items()):
        if qid in existing:
            existing_letter = existing[qid]
            if existing_letter is None or existing_letter == '':
                # Was NULL, now setting
                cur.execute(
                    "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
                    (answer_letter, qid)
                )
                new_count += 1
            elif existing_letter == answer_letter:
                same_count += 1
            else:
                # Override the wrong answer
                print(f"  OVERRIDE: {qid} '{existing_letter}' -> '{answer_letter}'")
                cur.execute(
                    "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
                    (answer_letter, qid)
                )
                conflict_count += 1
                update_count += 1
        else:
            cur.execute(
                "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
                (answer_letter, qid)
            )
            new_count += 1
    
    conn.commit()
    conn.close()
    
    # Final summary
    print(f"\n{'=' * 60}")
    print(f"FINAL SUMMARY")
    print(f"{'=' * 60}")
    print(f"  Total questions in source: 1119")
    print(f"  Total updates in mapping:  {len(all_updates)}")
    print(f"  Newly set (was NULL):      {new_count}")
    print(f"  Overridden (was wrong):    {conflict_count}")
    print(f"  Already correct:           {same_count}")
    
    # Check final state
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=3 AND correct_answer_letter IS NOT NULL AND correct_answer_letter != ''")
    filled = cur.fetchone()[0]
    cur.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=3")
    total = cur.fetchone()[0]
    conn.close()
    
    print(f"\n  After update:")
    print(f"    Questions WITH answer:    {filled}/{total} ({filled*100//total}%)")
    print(f"    Still missing:            {total-filled}")
    print(f"\n  Mapping saved to: {OUTPUT_PATH}")
    print(f"{'=' * 60}")


if __name__ == '__main__':
    main()
