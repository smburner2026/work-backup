#!/usr/bin/env python3
"""
Extract answers from ALL 'with answers' docx files and match to DB questions.
Version 2: Fixes text matching and handles more docx formats.
"""

import os, re, json, sqlite3, sys
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_JSON = "/root/work/dabt/kristen_mini_answers.json"

def extract_answers_v2(doc_path):
    """
    Improved answer extraction. Looks for bold text on answer option lines.
    Also checks for 'Answer:' or 'ANSWER:' markers.
    """
    doc = Document(doc_path)
    answers = {}
    current_q = None
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    m = re.match(r'^(\d+)\.\s+', text)
                    if m:
                        current_q = int(m.group(1))
                    
                    bold_texts = []
                    for run in para.runs:
                        if run.bold and run.text.strip():
                            bold_texts.append(run.text.strip())
                    
                    if bold_texts and current_q is not None:
                        full_bold = ''.join(bold_texts).strip()
                        opt_match = re.match(r'^([A-H])[\.\s\)]', full_bold)
                        if opt_match and current_q not in answers:
                            answers[current_q] = opt_match.group(1)
    
    # Main paragraph scan
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check for question number
        m = re.match(r'^(\d+)\.\s+', text)
        if m:
            current_q = int(m.group(1))
        
        # Check for bold runs
        bold_texts = []
        for run in para.runs:
            if run.bold and run.text.strip():
                bold_texts.append(run.text.strip())
        
        if bold_texts and current_q is not None:
            full_bold = ''.join(bold_texts).strip()
            # Match letter option like "A." or "A)" or "A" at start
            opt_match = re.match(r'^([A-H])[\.\s\)]', full_bold)
            if opt_match and current_q not in answers:
                answers[current_q] = opt_match.group(1)
    
    # If no answers found, try 'Answer:' approach
    if len(answers) == 0:
        for para in doc.paragraphs:
            text = para.text.strip()
            # Look for patterns like "Answer: A" or "ANSWER: B" or "1. A" 
            m = re.match(r'^(\d+)\.?\s*[:\-]?\s*([A-H])', text, re.IGNORECASE)
            if m:
                qnum = int(m.group(1))
                letter = m.group(2).upper()
                answers[qnum] = letter
            
            m2 = re.match(r'(?:Answer|ANSWER|ans\.?)\s*:\s*([A-H])', text)
            if m2:
                letter = m2.group(1).upper()
                # Associate with last seen question
                if current_q is not None and current_q not in answers:
                    answers[current_q] = letter
    
    return answers

def get_all_db_questions():
    """Fetch all questions from source_file_id=4."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("""
        SELECT id, question_number_in_source, question_text 
        FROM questions 
        WHERE source_file_id=4 
        ORDER BY id
    """)
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()
    return rows

def normalize_text(t):
    """Normalize text for comparison."""
    t = re.sub(r'\s+', ' ', t).strip().lower()
    t = re.sub(r'[^\w\s]', '', t)  # Remove punctuation
    return t

def main():
    # Get all "with answers" docx files
    all_files = sorted(os.listdir(EXAM_DIR))
    with_answers_files = [f for f in all_files if 'with answers' in f.lower() or 'examination with answers' in f.lower()]
    # Also include "exam answers" files
    answers_files = [f for f in all_files if 'answer' in f.lower() and f.endswith('.docx') and f not in with_answers_files]
    
    print(f"Found {len(with_answers_files)} 'with answers' files and {len(answers_files)} additional answer files")
    for f in with_answers_files + answers_files:
        print(f"  {f}")
    
    # Extract answers from each file
    all_exam_answers = {}
    for fname in with_answers_files + answers_files:
        fpath = os.path.join(EXAM_DIR, fname)
        answers = extract_answers_v2(fpath)
        print(f"\n{fname}: extracted {len(answers)} answers")
        if len(answers) > 0:
            # Show first few
            sample = dict(sorted(answers.items())[:5])
            print(f"  Sample: {sample}")
        all_exam_answers[fname] = answers
    
    # Get DB questions
    db_questions = get_all_db_questions()
    print(f"\nDB has {len(db_questions)} questions from source_file_id=4")
    
    # Strategy: Match by question text.
    # For each exam file, extract question texts from the docx.
    # Then match each exam's questions to DB questions by text similarity.
    
    # Build a mapping of question_text -> DB question for easy lookup
    # Normalize both sides
    
    # For each exam with answers, extract question texts
    exam_data = {}
    for fname, answers in all_exam_answers.items():
        if len(answers) == 0:
            continue
        
        fpath = os.path.join(EXAM_DIR, fname)
        doc = Document(fpath)
        
        qtexts = {}
        for para in doc.paragraphs:
            m = re.match(r'^(\d+)\.\s+(.*)', para.text.strip())
            if m:
                qnum = int(m.group(1))
                qtexts[qnum] = para.text.strip()
        
        exam_data[fname] = {
            'answers': answers,
            'qtexts': qtexts,
            'fpath': fpath
        }
    
    # Now match each exam to DB questions
    # Strategy: since each exam has unique question numbering (starting at 1),
    # and the DB has multiple questions per number, we match by text.
    
    # For each question in the DB, find which exam and which Q number it belongs to
    # by comparing text
    
    db_remaining = list(db_questions)
    id_to_answer = {}
    
    for fname, data in exam_data.items():
        answers = data['answers']
        qtexts = data['qtexts']
        
        print(f"\n{'='*60}")
        print(f"Processing: {fname}")
        print(f"  Answers: {len(answers)} questions, Qtexts: {len(qtexts)} questions")
        
        matched = 0
        
        for qnum in sorted(answers.keys()):
            answer_letter = answers[qnum]
            
            if qnum not in qtexts:
                continue
            
            exam_full_text = qtexts[qnum]
            # Remove the "N. " prefix for comparison
            exam_text = re.sub(r'^\d+\.\s+', '', exam_full_text).strip()
            exam_norm = normalize_text(exam_text)
            
            # Find best matching DB question
            best_match = None
            best_score = 0
            
            for db_q in db_remaining:
                if db_q['question_number_in_source'] != qnum:
                    continue
                
                db_text = db_q['question_text'].strip()
                db_norm = normalize_text(db_text)
                
                # Use simple prefix matching - find the longest common prefix
                score = 0
                for i in range(min(len(exam_norm), len(db_norm))):
                    if exam_norm[i] == db_norm[i]:
                        score += 1
                    else:
                        break
                
                # Also check if one contains the other
                if score > best_score:
                    best_score = score
                    best_match = db_q
            
            if best_match and best_score >= 15:
                id_to_answer[best_match['id']] = {
                    'answer_letter': answer_letter,
                    'exam': fname,
                    'qnum': qnum,
                    'match_score': best_score,
                    'exam_text': exam_text[:80],
                    'db_text': best_match['question_text'][:80],
                }
                db_remaining.remove(best_match)
                matched += 1
            else:
                if best_match:
                    print(f"  Q{qnum}: Low match score {best_score} for [{best_match['id']}]")
                    print(f"    Exam: {exam_text[:60]}")
                    print(f"    DB:   {best_match['question_text'][:60]}")
                else:
                    # Check if there are even candidates
                    candidates = [q for q in db_remaining if q['question_number_in_source'] == qnum]
                    if candidates:
                        print(f"  Q{qnum}: No good match (score={best_score}) among {len(candidates)} candidates, best: [{candidates[0]['id']}] {candidates[0]['question_text'][:60]}")
                    else:
                        print(f"  Q{qnum}: No DB candidates remaining (all already matched?)")
        
        print(f"  Matched {matched} questions in this exam")
        print(f"  Total matched so far: {len(id_to_answer)}")
    
    print(f"\n{'='*60}")
    print(f"Total matched: {len(id_to_answer)} / {len(db_questions)}")
    
    # Show unmatched
    unmatched = [q for q in db_questions if q['id'] not in id_to_answer]
    if unmatched:
        print(f"\nUnmatched DB questions ({len(unmatched)}):")
        for q in unmatched[:20]:
            print(f"  [{q['id']}] Q{q['question_number_in_source']}: {q['question_text'][:60]}")
    
    if len(unmatched) > 20:
        print(f"  ... and {len(unmatched)-20} more")
    
    # Save results as JSON
    result_json = {}
    for db_id, info in id_to_answer.items():
        result_json[db_id] = {
            'answer_letter': info['answer_letter'],
            'exam': info['exam'],
            'qnum': info['qnum']
        }
    
    with open(OUTPUT_JSON, 'w') as f:
        json.dump(result_json, f, indent=2)
    print(f"\nSaved to {OUTPUT_JSON}")
    
    # Update database
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    updated = 0
    errors = 0
    for db_id, info in result_json.items():
        try:
            cursor.execute(
                "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
                (info['answer_letter'], db_id)
            )
            if cursor.rowcount > 0:
                updated += 1
            else:
                print(f"  ERROR: No row updated for {db_id}")
                errors += 1
        except Exception as e:
            print(f"  ERROR updating {db_id}: {e}")
            errors += 1
    
    conn.commit()
    
    # Verify
    cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4 AND correct_answer_letter IS NOT NULL AND correct_answer_letter != ''")
    filled = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4")
    total = cursor.fetchone()[0]
    print(f"\nDB Update: {updated} updated, {errors} errors")
    print(f"After update: {filled}/{total} questions have correct_answer_letter")
    
    conn.close()

if __name__ == '__main__':
    main()
