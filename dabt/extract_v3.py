#!/usr/bin/env python3
"""
Final version: Extract answers from all 'with answers' docx files.
Handles multiple formats: "A. text", "A) text", "ANSWER: A", "Q1:" etc.
Matches to DB by question text similarity.
"""

import os, re, json, sqlite3, sys
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_JSON = "/root/work/dabt/kristen_mini_answers.json"

def extract_answers_comprehensive(doc_path):
    """Extract question_number -> answer_letter from any format docx."""
    doc = Document(doc_path)
    answers = {}
    current_q = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detect question number from various formats
        # Format 1: "N. text" or "N) text"
        m = re.match(r'^(\d+)[\.\)]\s+', text)
        if m:
            current_q = int(m.group(1))
        
        # Format 2: "QN:" or "QN." 
        m2 = re.match(r'^[Qq](\d+)[\.:]\s+', text)
        if m2:
            current_q = int(m2.group(1))
        
        # Check for bold runs
        bold_texts = []
        for run in para.runs:
            if run.bold and run.text.strip():
                bold_texts.append(run.text.strip())
        
        if not bold_texts:
            continue
        
        full_bold = ''.join(bold_texts).strip()
        
        # Format A: "ANSWER: X" or "ANSWER:X"
        m_ans = re.match(r'^(?:Answer|ANSWER)\s*:\s*([A-H])', full_bold)
        if m_ans and current_q is not None and current_q not in answers:
            answers[current_q] = m_ans.group(1)
            continue
        
        # Format B: "A. text" or "A) text" or "A text"
        # But ONLY if this looks like an answer option (starts with single letter)
        m_opt = re.match(r'^([A-H])[\.\)]\s', full_bold)
        if m_opt and current_q is not None and current_q not in answers:
            # Make sure it's not the question number line being bolded
            # Question lines look like "1. text" or "Q1: text"
            if not re.match(r'^\d+\.\s', full_bold) and not re.match(r'^[Qq]\d+', full_bold):
                answers[current_q] = m_opt.group(1)
                continue
        
        # Format C: "A, B, and C" type compound answers
        # These usually correspond to answer choices like G, H, etc.
        # We'll handle these separately - they're options like "G. A, B, and C"
        # where the entire "G. A, B, and C" is bolded
        m_comp = re.match(r'^([A-H])\.\s+([A-H])(?:,?\s*[a-zA-Z]+\s+)*([A-H]?)\s*', full_bold)
        # Actually, the simple case: bold starts with single letter followed by period or )
        m_simple = re.match(r'^([A-H])[\.\)]', full_bold)
        if m_simple and current_q is not None and current_q not in answers:
            # Additional check: don't match question number lines
            first_word = full_bold.split()[0] if full_bold.split() else ''
            if not first_word.lstrip('(').rstrip(')').isdigit():  # Not a number
                answers[current_q] = m_simple.group(1)
    
    return answers

def extract_question_texts(doc_path):
    """Extract question_number -> text from a docx."""
    doc = Document(doc_path)
    qtexts = {}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Format 1: "N. text"
        m = re.match(r'^(\d+)[\.\)]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            qtext = m.group(2).strip()
            qtexts[qnum] = qtext
    
    # Also check "QN:" format
    for para in doc.paragraphs:
        text = para.text.strip()
        m = re.match(r'^[Qq](\d+)[\.:]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            qtext = m.group(2).strip()
            if qnum not in qtexts:
                qtexts[qnum] = qtext
    
    return qtexts

def normalize(t):
    t = re.sub(r'\s+', ' ', t).strip().lower()
    t = re.sub(r'[^\w\s]', '', t)
    return t

def get_all_db_questions():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("""
        SELECT id, question_number_in_source, question_text 
        FROM questions WHERE source_file_id=4 
        ORDER BY id
    """)
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()
    return rows

def main():
    all_files = sorted(os.listdir(EXAM_DIR))
    
    # Find all "with answers" files
    with_files = [f for f in all_files if ('with answers' in f.lower() or 'examination with answers' in f.lower() or 'exam answers' in f.lower())]
    
    print(f"Found {len(with_files)} answer files")
    for f in with_files:
        print(f"  {f}")
    
    # Extract answers and question texts from each answer file
    exam_answers = {}
    exam_qtexts = {}
    
    for fname in with_files:
        fpath = os.path.join(EXAM_DIR, fname)
        answers = extract_answers_comprehensive(fpath)
        qtexts = extract_question_texts(fpath)
        exam_answers[fname] = answers
        exam_qtexts[fname] = qtexts
        
        print(f"\n{fname}:")
        print(f"  Answers: {len(answers)}")
        if len(answers) > 0:
            print(f"  Sample answers: {dict(sorted(answers.items())[:5])}")
        print(f"  Question texts: {len(qtexts)}")
    
    # Get DB questions
    db_questions = get_all_db_questions()
    print(f"\nDB has {len(db_questions)} questions from source_file_id=4")
    
    # Build normalized DB lookup by question_number_in_source
    db_by_qnum = {}
    for q in db_questions:
        qn = q['question_number_in_source']
        if qn not in db_by_qnum:
            db_by_qnum[qn] = []
        db_by_qnum[qn].append(q)
    
    # Now match each exam to DB questions
    # For each exam, iterate through its questions in order,
    # and match by text similarity with DB questions of the same question_number_in_source
    
    db_unmatched = set(q['id'] for q in db_questions)
    id_to_answer = {}
    
    for fname in with_files:
        answers = exam_answers[fname]
        qtexts = exam_qtexts[fname]
        
        if len(answers) == 0:
            continue
        
        print(f"\n{'='*60}")
        print(f"Matching: {fname}")
        matched_in_file = 0
        
        for qnum in sorted(answers.keys()):
            answer_letter = answers[qnum]
            
            if qnum not in qtexts:
                print(f"  Q{qnum}: No question text found in docx (ans={answer_letter}), skipping")
                continue
            
            exam_text = qtexts[qnum]
            exam_norm = normalize(exam_text)
            
            # Find candidates from DB with same qnum that are still unmatched
            candidates = [q for q in db_by_qnum.get(qnum, []) if q['id'] in db_unmatched]
            
            if len(candidates) == 0:
                # All candidates already matched - check if there ARE any
                all_candidates = db_by_qnum.get(qnum, [])
                if len(all_candidates) == 0:
                    print(f"  Q{qnum}: No DB candidates for this qnum (ans={answer_letter})")
                else:
                    print(f"  Q{qnum}: All {len(all_candidates)} DB candidates already matched (ans={answer_letter})")
                continue
            
            if len(candidates) == 1:
                # Only one unmatch candidate - use it
                best_match = candidates[0]
                best_score = 999
            else:
                # Multiple candidates - pick best text match
                best_match = None
                best_score = 0
                
                for c in candidates:
                    db_norm = normalize(c['question_text'])
                    
                    # Score: longest common prefix
                    score = 0
                    for i in range(min(len(exam_norm), len(db_norm))):
                        if exam_norm[i] == db_norm[i]:
                            score += 1
                        else:
                            break
                    
                    if score > best_score:
                        best_score = score
                        best_match = c
            
            if best_match:
                db_unmatched.discard(best_match['id'])
                id_to_answer[best_match['id']] = {
                    'answer_letter': answer_letter,
                    'exam': fname,
                    'qnum': qnum
                }
                matched_in_file += 1
                
                if len(candidates) > 1:
                    print(f"  Q{qnum}: Matched to [{best_match['id']}] (score={best_score}, from {len(candidates)} candidates)")
        
        print(f"  Matched {matched_in_file} questions in this file")
    
    print(f"\n{'='*60}")
    print(f"Total matched: {len(id_to_answer)} / {len(db_questions)}")
    
    unmatched_ids = [q['id'] for q in db_questions if q['id'] not in id_to_answer]
    print(f"Unmatched: {len(unmatched_ids)}")
    if unmatched_ids:
        for q in db_questions:
            if q['id'] in unmatched_ids:
                print(f"  [{q['id']}] Q{q['question_number_in_source']}: {q['question_text'][:60]}")
    
    # Save JSON
    result_json = {}
    for db_id, info in id_to_answer.items():
        result_json[db_id] = info
    
    with open(OUTPUT_JSON, 'w') as f:
        json.dump(result_json, f, indent=2)
    print(f"\nSaved to {OUTPUT_JSON}")
    
    # Update DB
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    updated = 0
    errors = 0
    for db_id, info in result_json.items():
        try:
            cursor.execute(
                "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
                (info['answer_letter'], db_id)
            )
            if cursor.rowcount > 0:
                updated += 1
            else:
                print(f"  ERROR: No row for {db_id}")
                errors += 1
        except Exception as e:
            print(f"  ERROR: {e}")
            errors += 1
    
    conn.commit()
    
    cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4 AND correct_answer_letter IS NOT NULL AND correct_answer_letter != ''")
    filled = cursor.fetchone()[0]
    total = len(db_questions)
    print(f"\nDB Updated: {updated} successful, {errors} errors")
    print(f"Questions with answers: {filled}/{total}")
    
    conn.close()

if __name__ == '__main__':
    main()
