#!/usr/bin/env python3
"""
FINAL VERSION: Match by TEXT SIMILARITY only, no sequential greedy matching.
For each DB question, find the best matching exam question across ALL docx files.
"""

import os, re, json, sqlite3
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_JSON = "/root/work/dabt/kristen_mini_answers.json"

def extract_answers(doc_path):
    """Extract question_number -> answer_letter from any format docx."""
    doc = Document(doc_path)
    answers = {}
    current_q = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detect question number: "N.", "N)", "QN:", "QN."
        m = re.match(r'^(\d+)[\.\)]\s+', text)
        if m:
            current_q = int(m.group(1))
        m2 = re.match(r'^[Qq](\d+)[\.:]\s+', text)
        if m2:
            current_q = int(m2.group(1))
        
        # Check for bold runs
        bold_texts = []
        for run in para.runs:
            if run.bold and run.text.strip():
                bold_texts.append(run.text.strip())
        
        if not bold_texts:
            continue
        
        full_bold = ''.join(bold_texts).strip()
        
        # "ANSWER: X" format
        m_ans = re.match(r'^(?:Answer|ANSWER)\s*:\s*([A-H])', full_bold)
        if m_ans and current_q is not None and current_q not in answers:
            answers[current_q] = m_ans.group(1)
            continue
        
        # "A. text" or "A) text" format - ONLY if it's an option line
        if current_q is not None and current_q not in answers:
            m_opt = re.match(r'^([A-H])[\.\)]', full_bold)
            if m_opt:
                # Skip question header lines
                if not re.match(r'^\d+\.\s', full_bold) and not re.match(r'^[Qq]\d+', full_bold):
                    answers[current_q] = m_opt.group(1)
    
    return answers

def extract_question_texts(doc_path):
    """Extract {qnum: clean_question_text} from docx."""
    doc = Document(doc_path)
    qtexts = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        m = re.match(r'^(\d+)[\.\)]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            qtext = m.group(2).strip()
            qtexts[qnum] = qtext
        m = re.match(r'^[Qq](\d+)[\.:]\s+(.*)', text)
        if m:
            qnum = int(m.group(1))
            qtext = m.group(2).strip()
            if qnum not in qtexts:
                qtexts[qnum] = qtext
    return qtexts

def normalize(t):
    t = re.sub(r'\s+', ' ', t).strip().lower()
    t = re.sub(r'[^\w\s]', '', t)
    return t

def text_similarity(a, b):
    """Character-by-character prefix match score."""
    an = normalize(a)
    bn = normalize(b)
    score = 0
    for i in range(min(len(an), len(bn))):
        if an[i] == bn[i]:
            score += 1
        else:
            break
    return score

# Main
all_files = sorted(os.listdir(EXAM_DIR))
answer_files = [f for f in all_files if ('with answers' in f.lower() or 'examination with answers' in f.lower() or 'exam answers' in f.lower())]

# For each answer file, extract answers and question texts
exam_data = {}
for fname in answer_files:
    fpath = os.path.join(EXAM_DIR, fname)
    answers = extract_answers(fpath)
    qtexts = extract_question_texts(fpath)
    exam_data[fname] = {'answers': answers, 'qtexts': qtexts, 'fpath': fpath}
    print(f"{fname}: {len(answers)} answers, {len(qtexts)} texts")

# Get DB questions
conn = sqlite3.connect(DB_PATH)
conn.row_factory = sqlite3.Row
cur = conn.cursor()
cur.execute("SELECT id, question_number_in_source, question_text FROM questions WHERE source_file_id=4 ORDER BY id")
db_questions = [dict(r) for r in cur.fetchall()]
conn.close()

print(f"\nDB has {len(db_questions)} questions")

# For each DB question, find the BEST matching exam question by text
id_to_answer = {}
unmatched = []

for db_q in db_questions:
    db_text = db_q['question_text']
    db_id = db_q['id']
    db_qnum = db_q['question_number_in_source']
    
    best_match = None
    best_score = 0
    
    for fname, data in exam_data.items():
        answers = data['answers']
        qtexts = data['qtexts']
        
        for qnum in answers.keys():
            if qnum not in qtexts:
                continue
            
            exam_text = qtexts[qnum]
            
            # Must have at least matching first few characters
            # Use question_number_in_source as soft filter
            if db_qnum != qnum:
                # Allow matching different qnums if text is very long match
                # But penalize by 10 to prefer same qnum
                score = text_similarity(exam_text, db_text) - 10
            else:
                score = text_similarity(exam_text, db_text)
            
            if score > best_score:
                best_score = score
                best_match = (fname, qnum, answers[qnum])
    
    if best_match and best_score >= 10:
        fname, qnum, answer = best_match
        id_to_answer[db_id] = {
            'answer_letter': answer,
            'exam': fname,
            'qnum': qnum,
            'match_score': best_score
        }
    else:
        unmatched.append((db_id, db_text, best_score, best_match))
        print(f"NO MATCH for [{db_id}] Q{db_qnum}: {db_text[:60]} (best_score={best_score})")

print(f"\nMatched: {len(id_to_answer)}/{len(db_questions)}")
print(f"Unmatched: {len(unmatched)}")

# Show any unmatched
for db_id, text, score, match in unmatched:
    print(f"  [{db_id}]: {text[:60]} (best_score={score})")

# Save JSON
with open(OUTPUT_JSON, 'w') as f:
    json.dump(id_to_answer, f, indent=2)
print(f"\nSaved to {OUTPUT_JSON}")

# Update DB
conn = sqlite3.connect(DB_PATH)
cursor = conn.cursor()
updated = 0
errors = 0
for db_id, info in id_to_answer.items():
    try:
        cursor.execute(
            "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
            (info['answer_letter'], db_id)
        )
        if cursor.rowcount > 0:
            updated += 1
        else:
            errors += 1
    except Exception as e:
        print(f"ERROR: {e}")
        errors += 1
conn.commit()

cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4 AND correct_answer_letter IS NOT NULL AND correct_answer_letter != ''")
filled = cursor.fetchone()[0]
total = len(db_questions)
print(f"\nDB Updated: {updated} ok, {errors} errors")
print(f"Questions with answers: {filled}/{total}")
conn.close()
