#!/usr/bin/env python3
"""Extract answers from ALL 'with answers' docx files and match to DB questions."""

import os, re, json, sqlite3, sys
from docx import Document

EXAM_DIR = "/root/dabt-curated/Practice_Exams/Kristen_Mini_Exams"
DB_PATH = "/root/work/dabt/dabt-tutor/reference/data/dabt.db"
OUTPUT_JSON = "/root/work/dabt/kristen_mini_answers.json"

def extract_answers_from_docx(doc_path):
    """
    Extract question_number -> answer_letter from a 'with answers' docx.
    The correct answer is the bolded option letter (e.g., "A.", "B.", etc.)
    """
    doc = Document(doc_path)
    answers = {}
    current_q = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this is a question number line (e.g., "1. What is...")
        m = re.match(r'^(\d+)\.\s+', text)
        if m:
            current_q = int(m.group(1))
        
        # Check for bold runs that indicate the answer
        bold_texts = []
        for run in para.runs:
            if run.bold and run.text.strip():
                bold_texts.append(run.text.strip())
        
        if bold_texts and current_q is not None:
            full_bold = ''.join(bold_texts).strip()
            # Match letter option like "A." or "A" at start, possibly with text after
            # Handle "A. text", "A text", "A, B, and C", "A, C, and F", etc.
            # Also handle compound answers like "A, B, and C" -> "G" in multiple choice
            opt_match = re.match(r'^([A-H])[\.\s]', full_bold)
            if opt_match:
                letter = opt_match.group(1)
                answers[current_q] = letter
            else:
                # Also try matching "A, B" type compound answers
                compound_match = re.match(r'^([A-H]),?\s*([A-H])', full_bold)
                if compound_match:
                    # These are compound answer choices like "A, B" which might be option "G"
                    # Just take first letter for now - these need special handling
                    answers[current_q] = full_bold[:2]  # placeholder
    
    return answers

def get_all_db_questions():
    """Fetch all questions from source_file_id=4."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("""
        SELECT id, question_number_in_source, question_text 
        FROM questions 
        WHERE source_file_id=4 
        ORDER BY id
    """)
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()
    return rows

def match_questions_to_exam(db_questions, exam_answers, exam_name):
    """
    Match DB questions to exam answers by comparing question text.
    Returns list of (db_id, question_number_in_source, answer_letter) tuples.
    """
    matches = []
    
    # For each answer in the exam, try to find the matching DB question
    for qnum, answer_letter in exam_answers.items():
        # Find DB questions with this question_number_in_source
        candidates = [q for q in db_questions if q['question_number_in_source'] == qnum]
        
        if len(candidates) == 0:
            print(f"  WARNING: No DB questions found for Q{qnum} in {exam_name}")
            continue
        elif len(candidates) == 1:
            matches.append((candidates[0]['id'], qnum, answer_letter))
            candidates[0]['_matched'] = True
        else:
            # Multiple candidates - need to match by context (position within exam)
            # Since we process exams in order of IDs, the first unmatch is likely the next one
            print(f"  WARNING: {len(candidates)} DB questions for Q{qnum} in {exam_name}")
            # We'll handle this by position
            for c in candidates:
                if not c.get('_matched'):
                    matches.append((c['id'], qnum, answer_letter))
                    c['_matched'] = True
                    break
    
    return matches

def main():
    # Get all "with answers" docx files
    all_files = sorted(os.listdir(EXAM_DIR))
    with_answers_files = [f for f in all_files if 'with answers' in f.lower() or 'examination with answers' in f.lower()]
    
    print(f"Found {len(with_answers_files)} 'with answers' files:")
    for f in with_answers_files:
        print(f"  {f}")
    
    # Extract answers from each file
    all_exam_answers = {}
    for fname in with_answers_files:
        fpath = os.path.join(EXAM_DIR, fname)
        answers = extract_answers_from_docx(fpath)
        print(f"\n{fname}: extracted {len(answers)} answers")
        all_exam_answers[fname] = answers
    
    # Get DB questions
    db_questions = get_all_db_questions()
    print(f"\nDB has {len(db_questions)} questions from source_file_id=4")
    
    # Show question_number_in_source distribution
    qnum_counts = {}
    for q in db_questions:
        qn = q['question_number_in_source']
        qnum_counts[qn] = qnum_counts.get(qn, 0) + 1
    print(f"Question number distribution: {dict(sorted(qnum_counts.items()))}")
    
    # Now try to match. The DB questions are ordered by ID.
    # Let's figure out which exam corresponds to which range of DB IDs.
    
    # Strategy: For each exam, extract its answers and try to match by question text
    # Since we have the full question text in the docx, let's compare directly
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # First, let's look at the first few questions of 26 August 2017 to match
    sample_exam = "Mini-ABT exam with answers 26 August 2017.docx"
    sample_answers = all_exam_answers[sample_exam]
    
    print(f"\n=== Attempting text-based matching for {sample_exam} ===")
    
    # Get full text of first question in 26 Aug 2017 exam
    doc = Document(os.path.join(EXAM_DIR, sample_exam))
    for i, para in enumerate(doc.paragraphs[:2]):
        print(f"  Para {i}: {para.text[:100]}")
    
    # Let me examine the actual question texts more carefully
    for qnum in [1, 2, 3, 4, 5]:
        answer = sample_answers.get(qnum)
        # Find the question text from the docx
        qtext_docx = ""
        doc = Document(os.path.join(EXAM_DIR, sample_exam))
        for para in doc.paragraphs:
            m = re.match(rf'^{qnum}\.\s+(.*)', para.text.strip())
            if m:
                qtext_docx = m.group(1)
                break
        
        # Find matching DB questions
        db_matches = [q for q in db_questions if q['question_number_in_source'] == qnum]
        print(f"\n  Q{qnum} (answer: {answer}):")
        print(f"    Docx: {qtext_docx[:80]}...")
        for db_q in db_matches:
            print(f"    DB:   [{db_q['id']}] {db_q['question_text'][:80]}...")
    
    # The approach: match by question text similarity. Let me build a text matcher.
    
    # Extract question texts from docx files
    print("\n\n=== Building text matching map ===")
    
    # For each exam, extract question texts
    exam_question_texts = {}
    for fname in with_answers_files:
        fpath = os.path.join(EXAM_DIR, fname)
        doc = Document(fpath)
        qtexts = {}
        for para in doc.paragraphs:
            m = re.match(r'^(\d+)\.\s+(.*)', para.text.strip())
            if m:
                qnum = int(m.group(1))
                qtexts[qnum] = para.text.strip()
        exam_question_texts[fname] = qtexts
    
    # Now for each DB question, find the best matching exam+question
    db_questions_remaining = list(db_questions)
    id_to_answer = {}
    
    for fname in with_answers_files:
        qtexts = exam_question_texts[fname]
        fpath = os.path.join(EXAM_DIR, fname)
        answers = all_exam_answers[fname]
        
        print(f"\nProcessing {fname}...")
        matched_in_exam = 0
        
        for qnum, answer_letter in sorted(answers.items()):
            if qnum not in qtexts:
                continue
            
            exam_text = qtexts[qnum]
            # Normalize: remove whitespace, lowercase
            exam_text_norm = re.sub(r'\s+', ' ', exam_text).strip().lower()
            
            # Find best matching DB question
            best_match = None
            best_score = 0
            
            for db_q in db_questions_remaining:
                if db_q['question_number_in_source'] != qnum:
                    continue
                
                db_text = db_q['question_text'].strip()
                db_text_norm = re.sub(r'\s+', ' ', db_text).lower()
                
                # Simple score: length of common prefix
                score = 0
                for i in range(min(len(exam_text_norm), len(db_text_norm))):
                    if exam_text_norm[i] == db_text_norm[i]:
                        score += 1
                    else:
                        break
                
                if score > best_score:
                    best_score = score
                    best_match = db_q
            
            if best_match and best_score > 20:
                id_to_answer[best_match['id']] = {
                    'answer_letter': answer_letter,
                    'exam': fname,
                    'qnum': qnum,
                    'match_score': best_score
                }
                db_questions_remaining.remove(best_match)
                matched_in_exam += 1
            else:
                print(f"  Could not match Q{qnum} (ans={answer_letter}), best_score={best_score}")
                if best_match:
                    print(f"    Best DB: [{best_match['id']}] {best_match['question_text'][:60]}")
                else:
                    print(f"    No DB candidates for Q{qnum}")
        
        print(f"  Matched {matched_in_exam}/{len(answers)} questions")
    
    print(f"\nTotal matched: {len(id_to_answer)} / {len(db_questions)}")
    
    # Check unmatched
    unmatched = [q for q in db_questions if q['id'] not in id_to_answer]
    if unmatched:
        print(f"\nUnmatched DB questions ({len(unmatched)}):")
        for q in unmatched[:20]:
            print(f"  [{q['id']}] Q{q['question_number_in_source']}: {q['question_text'][:60]}")
    
    # Save results as JSON
    result_json = {}
    for db_id, info in id_to_answer.items():
        result_json[db_id] = info
    
    with open(OUTPUT_JSON, 'w') as f:
        json.dump(result_json, f, indent=2)
    print(f"\nSaved results to {OUTPUT_JSON}")
    
    # Update database
    updated = 0
    errors = 0
    for db_id, info in id_to_answer.items():
        try:
            cursor.execute(
                "UPDATE questions SET correct_answer_letter = ? WHERE id = ?",
                (info['answer_letter'], db_id)
            )
            if cursor.rowcount > 0:
                updated += 1
            else:
                print(f"  ERROR: No update for {db_id}")
                errors += 1
        except Exception as e:
            print(f"  ERROR updating {db_id}: {e}")
            errors += 1
    
    conn.commit()
    print(f"\nDB Update: {updated} updated, {errors} errors")
    
    # Verify
    cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4 AND correct_answer_letter IS NOT NULL AND correct_answer_letter != ''")
    filled = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT(*) FROM questions WHERE source_file_id=4")
    total = cursor.fetchone()[0]
    print(f"After update: {filled}/{total} questions have correct_answer_letter")
    
    conn.close()

if __name__ == '__main__':
    main()
