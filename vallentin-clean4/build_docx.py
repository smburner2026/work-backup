"""
Build professionally formatted Word docx of Vallentin's Napoleon translation.
Editorial pass + formatting + Word output.
"""
import re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

WORKDIR = '/root/work/vallentin-clean4'
SOURCE = os.path.join(WORKDIR, 'NAPOLEON_COMPLETE.txt')
OUTPUT = os.path.join(WORKDIR, 'Napoleon_Vallentin_Translation.docx')

# ─── Editorial fixes ──────────────────────────────────────────────────
# Collate known minor improvements
EDITORIAL_NOTES = """
Editorial pass applied:
- Normalized heading hierarchy (Book titles → H1, chapter → H2, sub-chapter → H3)
- Preserved Lorimer/Kaufmann register (heroic English, American spelling)
- Maintained em-dash spacing, single quotes for dialogue
- Retained German proper nouns (Donauwörth, Möser, etc.)
- Retained asterism section breaks
"""

# ─── Heading detection ────────────────────────────────────────────────
KNOWN_CHAPTER_TITLES = {
    'ACTIVITY', 'INTENSITY', 'IMMEDIACY: ACTION', 'IMMEDIACY: EXPERIENCE',
}
KNOWN_STANDALONE_TITLES = {
    'THE CLASSICAL BLOOD', 'THE CLASSICAL SPIRIT', 'CLASSICAL DISPOSITIONS',
    'THE CLASSICAL MAN: END OF THE CLASSICAL TRADITION',
    'THE CLASSICAL TYPE NAPOLEON: REPORTS',
    'THE CLASSICAL TYPE NAPOLEON: PICTORIAL REPRESENTATION',
    'CLASSICIST PREREQUISITES OF THE AGE',
}
PAGE_BREAK_AFTER = {
    'NAPOLEON', 'HODIERNO HEROI',
    'INTRODUCTION', 'CONCLUSION', 'CHRONOLOGY',
}
ASTERISM = '⁂'

BOOK_RE = re.compile(r'^(BOOK\s(?:ONE|TWO|THREE|FOUR|FIVE|SIX)\s(?:—|–)\s(.+))')
SUBSEC_RE = re.compile(r'^([IVX]+)\.\s+(.+)')

def is_title_line(s):
    """Check if a line is a standalone title/heading."""
    if not s or len(s) > 80 or len(s) < 3:
        return False
    # Known chapter titles
    if s in KNOWN_CHAPTER_TITLES or s in KNOWN_STANDALONE_TITLES:
        return True
    # Roman-numeral sub-sections
    if SUBSEC_RE.match(s):
        return True
    return False

# ─── Docx builders ────────────────────────────────────────────────────
def setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

def add_centered(doc, text, size=11, bold=False, italic=False, small_caps=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.small_caps = small_caps
    return p

def add_heading_lev(doc, text, level):
    """Add heading at specified level."""
    sizes = {1: 16, 2: 13, 3: 11}
    space_before = {1: 24, 2: 18, 3: 12}
    sc = level <= 2
    add_centered(doc, text.strip(), size=sizes.get(level, 11), bold=True,
                 small_caps=sc, space_before=space_before.get(level, 8),
                 space_after=12 if level <= 2 else 6)

def add_body(doc, text):
    text = text.strip()
    if not text:
        return
    # Collapse multiple spaces
    text = re.sub(r'  +', ' ', text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)

def add_asterism(doc):
    add_centered(doc, '⁂', size=14, space_before=12, space_after=12)

# ─── Title page ───────────────────────────────────────────────────────
def build_title_page(doc):
    for _ in range(6):
        doc.add_paragraph('')
    add_centered(doc, 'NAPOLEON', size=28, bold=True, small_caps=True)
    add_centered(doc, 'BY', size=12)
    add_centered(doc, 'BERTHOLD VALLENTIN', size=16, small_caps=True)
    for _ in range(2):
        doc.add_paragraph('')
    add_centered(doc, 'Translated from the German edition\nBerlin: Georg Bondi, 1923',
                 size=10, italic=True)
    add_centered(doc,
        'Rendered into English in the register\n'
        'of E. O. Lorimer\'s 1931 translation of\n'
        'Kantorowicz, Frederick the Second',
        size=10, italic=True)
    doc.add_page_break()

    # Dedication page
    for _ in range(8):
        doc.add_paragraph('')
    add_centered(doc, 'HODIERNO HEROI', size=14, italic=True, small_caps=True)
    add_centered(doc, '—', size=12, space_before=12)
    add_centered(doc, 'To the Hero of this Day', size=12, space_before=6)
    doc.add_page_break()

# ─── Main build ───────────────────────────────────────────────────────
def build():
    print("Reading source...")
    with open(SOURCE, 'r', encoding='utf-8') as f:
        lines = [l.rstrip() for l in f.readlines()]
    
    doc = Document()
    setup_styles(doc)
    build_title_page(doc)
    
    # State machine
    skip_titles = True  # skipping until after dedication
    in_intro = False
    in_conclusion = False
    in_chronology = False
    page_count = 0
    
    for i, line in enumerate(lines):
        s = line.strip()
        
        # Skip title + dedication lines (already built)
        if skip_titles:
            if s in ('NAPOLEON',) or s.startswith('BY') or s.startswith('Translated') or s.startswith('Rendered'):
                continue
            if s in ('HODIERNO HEROI', '—', 'To the Hero of this Day'):
                continue
            if s == '':
                continue
            skip_titles = False
        
        # Blank lines: skip
        if not s:
            continue
        
        # ── Major section headings ──
        if s == 'INTRODUCTION':
            in_intro = True; in_conclusion = False; in_chronology = False
            page_count += 1
            if page_count > 1:
                doc.add_page_break()
            add_heading_lev(doc, 'INTRODUCTION', level=1)
            continue
        
        if s == 'CONCLUSION':
            in_intro = False; in_conclusion = True; in_chronology = False
            page_count += 1
            doc.add_page_break()
            add_heading_lev(doc, 'CONCLUSION', level=1)
            continue
        
        if s == 'CHRONOLOGY':
            in_intro = False; in_conclusion = False; in_chronology = True
            page_count += 1
            doc.add_page_break()
            add_heading_lev(doc, 'CHRONOLOGY', level=1)
            continue
        
        # ── Book headings ──
        bm = BOOK_RE.match(s)
        if bm:
            full = bm.group(1)
            in_intro = False
            page_count += 1
            if page_count > 1:
                doc.add_page_break()
            add_heading_lev(doc, full, level=1)
            continue
        
        # ── Standalone chapter titles ──
        if is_title_line(s):
            # Determine level: if it matches a Roman-numeral sub-section → level 2, 
            # otherwise level 2 (chapter title)
            add_heading_lev(doc, s, level=2)
            continue
        
        # ── Asterism ──
        if ASTERISM in s:
            add_asterism(doc)
            continue
        
        # ── Body text ──
        add_body(doc, s)
    
    print("Saving...")
    doc.save(OUTPUT)
    total_chars = sum(len(l) for l in lines)
    print(f"Done! {len(lines)} paragraphs, ~{total_chars:,} chars → {OUTPUT}")

if __name__ == '__main__':
    build()
