"""
Build professionally formatted Word docx from individual chapter files
in correct order. Does editorial pass + formatting.
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

WORKDIR = '/root/work/vallentin-clean4'
OUTPUT = os.path.join(WORKDIR, 'Napoleon_Vallentin_Translation.docx')

# ─── Chapter order (explicit, correct sequence) ───────────────────────
# Each entry: (filename, display_book_title_or_None)
# If display_book_title is set, it's shown as a book-level heading.
CHAPTERS = [
    # Dedication
    ('dedication.en.txt', None),
    # Introduction
    ('00_einleitung.en.txt', 'INTRODUCTION'),
    # Book One: Deed and Experience
    ('01_aktivität.en.txt', 'BOOK ONE — DEED AND EXPERIENCE'),
    ('01_intensität.en.txt', None),
    ('01_unmittelbarkeit_handeln.en.txt', None),
    ('01_unmittelbarkeit_erleben.en.txt', None),
    # Book Two: History and the Present
    ('02_das_erlebnis_der_geschichte.en.txt', 'BOOK TWO — HISTORY AND THE PRESENT'),
    ('02_wandlungen_des_geschichtlichen_erlebens.en.txt', None),
    ('02_geschichte_und_tat.en.txt', None),
    # Book Three: The Classical
    ('03_klassische_anlagen.en.txt', 'BOOK THREE — THE CLASSICAL'),
    ('03_klassizistische_zeitvoraussetzungen.en.txt', None),
    ('03_das_klassische_geblüt.en.txt', None),
    ('03_der_klassische_geist.en.txt', None),
    ('03_der_klassische_typus_napoleon_bildliche_darstellung.en.txt', None),
    ('03_der_klassische_typus_napoleon_berichte.en.txt', None),
    ('03_der_klassische_mensch_ende_der_klassischen_tradition.en.txt', None),
    # Book Four: Feelings and Drives
    ('04_zeitströmungen.en.txt', 'BOOK FOUR — FEELINGS AND DRIVES'),
    ('04_eigenes_fühlen_unmittelbarkeit.en.txt', None),
    ('04_sendung_selbstgefühl_staatliches_empfinden.en.txt', None),
    # Book Five: God and Faith
    ('05_gott.en.txt', 'BOOK FIVE — GOD AND FAITH'),
    ('05_konfession.en.txt', None),
    ('05_kirche_ritus.en.txt', None),
    ('05_mythus_und_dogma_jugendperiode.en.txt', None),
    ('05_mythus_und_dogma_reifezeit.en.txt', None),
    ('05_glauben_und_staat.en.txt', None),
    # Book Six: Art
    ('06_dichtung_anlagen_und_betätigung.en.txt', 'BOOK SIX — ART'),
    ('06_dichterische_neigungen.en.txt', None),
    ('06_die_tragödie.en.txt', None),
    ('06_tragödie_und_mythus.en.txt', None),
    ('06_dichtung_und_politik.en.txt', None),
    ('06_bildende_kunst.en.txt', None),
    ('06_baukunst.en.txt', None),
    ('06_musik.en.txt', None),
    ('06_bildnisse.en.txt', None),
    # Conclusion & Chronology
    ('07_schlusswort.en.txt', 'CONCLUSION'),
    ('08_zeittafel.en.txt', 'CHRONOLOGY'),
]

# Roman numeral to Arabic for ordering within books
ROMAN_NUM = {
    'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
    'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9
}

ASTERISM = '⁂'

# ─── Editorial regex fixes ──────────────────────────────────────────
EDITORIAL_FIXES = [
    (r'  +', ' '),  # collapse double spaces
]

# ─── Heading parsing ─────────────────────────────────────────────────
BOOK_HEADING_RE = re.compile(
    r'^BOOK (ONE|TWO|THREE|FOUR|FIVE|SIX) — (.+)'
)
SUBSEC_HEADING_RE = re.compile(
    r'^([IVX]+)\.\s+(.+)'
)

def extract_heading(line):
    """Parse heading line and return (level, clean_text) or None."""
    s = line.strip()
    if not s:
        return None
    
    # Major sections
    if s == 'INTRODUCTION':
        return (1, 'INTRODUCTION')
    if s == 'CONCLUSION':
        return (1, 'CONCLUSION')
    if s == 'CHRONOLOGY':
        return (1, 'CHRONOLOGY')
    
    # Book headings
    bm = BOOK_HEADING_RE.match(s)
    if bm:
        return (1, s)
    
    # Sub-section: "I. GOD", "VI. FAITH AND THE STATE" etc
    sm = SUBSEC_HEADING_RE.match(s)
    if sm:
        return (2, s)
    
    # Known standalone chapter titles (within a book)
    chapter_titles = {
        'ACTIVITY', 'INTENSITY', 
        'IMMEDIACY: ACTION', 'IMMEDIACY: EXPERIENCE',
        'THE EXPERIENCE OF HISTORY',
        'TRANSFORMATIONS OF HISTORICAL EXPERIENCE',
        'HISTORY AND DEED',
        'CLASSICAL DISPOSITIONS', 'CLASSICIST PREREQUISITES OF THE AGE',
        'THE CLASSICAL BLOOD', 'THE CLASSICAL SPIRIT',
        'THE CLASSICAL TYPE NAPOLEON: PICTORIAL REPRESENTATION',
        'THE CLASSICAL TYPE NAPOLEON: REPORTS',
        'THE CLASSICAL MAN: END OF THE CLASSICAL TRADITION',
        "ONE'S OWN FEELING: IMMEDIACY",
        'MISSION. SELF-FEELING. POLITICAL SENTIMENT',
        'CURRENTS OF THE AGE',
        'MYTH AND DOGMA: YOUTH PERIOD', 'MYTH AND DOGMA: MATURITY',
        'FAITH AND THE STATE', 'CHURCH. RITE', 'CONFESSION', 'GOD',
        'POETRY: DISPOSITIONS AND PRACTICE',
        'POETIC INCLINATIONS', 'TRAGEDY',
        'TRAGEDY AND MYTH', 'POETRY AND POLITICS',
        'VISUAL ARTS', 'ARCHITECTURE', 'MUSIC', 'PORTRAITS',
    }
    if s in chapter_titles:
        return (2, s)
    
    return None

def apply_editorial(text):
    """Apply editorial fixes to body text."""
    for pattern, replacement in EDITORIAL_FIXES:
        text = re.sub(pattern, replacement, text)
    return text

# ─── Docx styling ────────────────────────────────────────────────────
def setup_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

def add_centered_p(doc, text, size=11, bold=False, italic=False, sc=False, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if sc:
        run.font.small_caps = True

def add_heading(doc, text, level):
    sizes = {1: 16, 2: 13}
    sb = {1: 30, 2: 20}
    sc = level == 1
    add_centered_p(doc, text.strip(), size=sizes.get(level, 11), bold=True,
                   sc=sc, sb=sb.get(level, 8), sa=12)

def add_body_para(doc, text):
    text = apply_editorial(text.strip())
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)

def add_asterism(doc):
    add_centered_p(doc, '⁂', size=14, sb=12, sa=12)

# ─── Title page ──────────────────────────────────────────────────────
def build_title_page(doc):
    for _ in range(6):
        doc.add_paragraph('')
    add_centered_p(doc, 'NAPOLEON', size=28, bold=True, sc=True)
    add_centered_p(doc, 'BY', size=12)
    add_centered_p(doc, 'BERTHOLD VALLENTIN', size=16, sc=True)
    for _ in range(2):
        doc.add_paragraph('')
    add_centered_p(doc, 'Translated from the German edition\nBerlin: Georg Bondi, 1923', size=10, italic=True)
    add_centered_p(doc,
        'Rendered into English in the register\n'
        'of E. O. Lorimer\'s 1931 translation of\n'
        'Kantorowicz, Frederick the Second',
        size=10, italic=True)
    doc.add_page_break()
    # Dedication
    for _ in range(8):
        doc.add_paragraph('')
    add_centered_p(doc, 'HODIERNO HEROI', size=14, italic=True, sc=True)
    add_centered_p(doc, '—', size=12, sb=12)
    add_centered_p(doc, 'To the Hero of this Day', size=12, sb=6)
    doc.add_page_break()

# ─── Build ───────────────────────────────────────────────────────────
def build():
    doc = Document()
    setup_doc(doc)
    build_title_page(doc)
    
    book_page_count = 0
    last_book = None
    
    for fname, book_title in CHAPTERS:
        path = os.path.join(WORKDIR, fname)
        with open(path, 'r', encoding='utf-8') as f:
            raw = f.read()
        
        lines = raw.split('\n')
        
        # Filter out empty trailing/leading lines
        content_lines = [l.rstrip() for l in lines]
        
        # Determine if this file starts a new book (page break before it)
        # Except for the introduction and dedication
        if book_title == 'INTRODUCTION':
            add_heading(doc, 'INTRODUCTION', level=1)
            book_page_count += 1
        elif book_title == 'CONCLUSION':
            book_page_count += 1
            doc.add_page_break()
            add_heading(doc, 'CONCLUSION', level=1)
        elif book_title == 'CHRONOLOGY':
            book_page_count += 1
            doc.add_page_break()
            add_heading(doc, 'CHRONOLOGY', level=1)
        elif book_title and book_title.startswith('BOOK '):
            if book_title != last_book:
                book_page_count += 1
                if book_page_count > 1:
                    doc.add_page_break()
                add_heading(doc, book_title, level=1)
                last_book = book_title
        elif book_title is None and fname != 'dedication.en.txt':
            # Continuation chapter within same book
            doc.add_page_break()
        
        # Process content lines
        for i, line in enumerate(content_lines):
            s = line.strip()
            if not s:
                continue
            
            # Skip the heading line if it matches the main heading patterns
            # (We've already added the book heading above)
            if i == 0:  # First line is usually the heading
                h = extract_heading(s)
                if h:
                    level, text = h
                    # For chapter-level headings within a book
                    if level == 2:
                        add_heading(doc, text, level=2)
                    # For book-level headings (already printed above)
                    elif level == 1:
                        # Skip - already added via book_title
                        pass
                    continue
            
            # Check for asterism
            if ASTERISM in s:
                add_asterism(doc)
                continue
            
            # Check if it's a heading appearing mid-file
            h = extract_heading(s)
            if h:
                level, text = h
                add_heading(doc, text, level=level)
                continue
            
            # Body text
            add_body_para(doc, s)
    
    doc.save(OUTPUT)
    print(f"Done! → {OUTPUT}")
    
    # Stats
    total_chars = sum(os.path.getsize(os.path.join(WORKDIR, f[0])) for f in CHAPTERS)
    print(f"Total input: {total_chars:,} bytes")
    sz = os.path.getsize(OUTPUT)
    print(f"Output: {sz:,} bytes ({sz/1024:.0f} KB)")

if __name__ == '__main__':
    build()
