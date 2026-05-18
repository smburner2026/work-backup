"""
Final build: professionally formatted Word docx of Vallentin's Napoleon.
- Correct chapter ordering
- Clean heading hierarchy
- Editorial pass applied
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

WORKDIR = '/root/work/vallentin-clean4'
OUTPUT = os.path.join(WORKDIR, 'Napoleon_Vallentin_Translation.docx')

# ─── Chapter definitions ──────────────────────────────────────────────
# Each: (filename, label_for_page_break, chapter_title_override)
# label_for_page_break: used before file; if starts with 'BOOK ' → page break
# chapter_title_override: if set, force this as the chapter heading (or None to auto-detect)

CHAPTERS = [
    ('dedication.en.txt',            None,                None),
    ('00_einleitung.en.txt',         'INTRODUCTION',      None),
    ('01_aktivität.en.txt',          'BOOK ONE — DEED AND EXPERIENCE',  'I. ACTIVITY'),
    ('01_intensität.en.txt',         None,                               'II. INTENSITY'),
    ('01_unmittelbarkeit_handeln.en.txt', None,                         'III. IMMEDIACY: ACTION'),
    ('01_unmittelbarkeit_erleben.en.txt', None,                         'IV. IMMEDIACY: EXPERIENCE'),
    ('02_das_erlebnis_der_geschichte.en.txt',  'BOOK TWO — HISTORY AND THE PRESENT',  'I. THE EXPERIENCE OF HISTORY'),
    ('02_wandlungen_des_geschichtlichen_erlebens.en.txt', None,         'II. TRANSFORMATIONS OF HISTORICAL EXPERIENCE'),
    ('02_geschichte_und_tat.en.txt',  None,                             'III. HISTORY AND DEED'),
    ('03_klassische_anlagen.en.txt', 'BOOK THREE — THE CLASSICAL',      'I. CLASSICAL DISPOSITIONS'),
    ('03_klassizistische_zeitvoraussetzungen.en.txt', None,             'II. CLASSICIST PREREQUISITES OF THE AGE'),
    ('03_das_klassische_geblüt.en.txt', None,                           'III. THE CLASSICAL BLOOD'),
    ('03_der_klassische_geist.en.txt', None,                            'IV. THE CLASSICAL SPIRIT'),
    ('03_der_klassische_typus_napoleon_bildliche_darstellung.en.txt', None, 'V. THE CLASSICAL TYPE NAPOLEON: PICTORIAL REPRESENTATION'),
    ('03_der_klassische_typus_napoleon_berichte.en.txt', None,          'VI. THE CLASSICAL TYPE NAPOLEON: REPORTS'),
    ('03_der_klassische_mensch_ende_der_klassischen_tradition.en.txt', None, 'VII. THE CLASSICAL MAN: END OF THE CLASSICAL TRADITION'),
    ('04_zeitströmungen.en.txt',     'BOOK FOUR — FEELINGS AND DRIVES', 'I. CURRENTS OF THE AGE'),
    ('04_eigenes_fühlen_unmittelbarkeit.en.txt', None,                  "II. ONE'S OWN FEELING: IMMEDIACY"),
    ('04_sendung_selbstgefühl_staatliches_empfinden.en.txt', None,      'III. MISSION. SELF-FEELING. POLITICAL SENTIMENT'),
    ('05_gott.en.txt',               'BOOK FIVE — GOD AND FAITH',       'I. GOD'),
    ('05_konfession.en.txt',         None,                              'II. CONFESSION'),
    ('05_kirche_ritus.en.txt',       None,                              'III. CHURCH. RITE'),
    ('05_mythus_und_dogma_jugendperiode.en.txt', None,                  'IV. MYTH AND DOGMA: YOUTH PERIOD'),
    ('05_mythus_und_dogma_reifezeit.en.txt', None,                      'V. MYTH AND DOGMA: MATURITY'),
    ('05_glauben_und_staat.en.txt',  None,                              'VI. FAITH AND THE STATE'),
    ('06_dichtung_anlagen_und_betätigung.en.txt', 'BOOK SIX — ART',     'I. POETRY: DISPOSITIONS AND PRACTICE'),
    ('06_dichterische_neigungen.en.txt', None,                          'II. POETIC INCLINATIONS'),
    ('06_die_tragödie.en.txt',       None,                              'III. TRAGEDY'),
    ('06_tragödie_und_mythus.en.txt', None,                             'IV. TRAGEDY AND MYTH'),
    ('06_dichtung_und_politik.en.txt', None,                            'V. POETRY AND POLITICS'),
    ('06_bildende_kunst.en.txt',     None,                              'VI. VISUAL ARTS'),
    ('06_baukunst.en.txt',           None,                              'VII. ARCHITECTURE'),
    ('06_musik.en.txt',              None,                              'VIII. MUSIC'),
    ('06_bildnisse.en.txt',          None,                              'IX. PORTRAITS'),
    ('07_schlusswort.en.txt',        'CONCLUSION',         None),
    ('08_zeittafel.en.txt',          'CHRONOLOGY',         None),
]

ASTERISM = '⁂'

# ─── Editorial fixes ──────────────────────────────────────────────────
EDITORIAL = [
    (r'  +', ' '),
]

def fix_text(text):
    for pat, repl in EDITORIAL:
        text = re.sub(pat, repl, text)
    return text

# ─── Docx styling ────────────────────────────────────────────────────
def setup(doc):
    s = doc.styles['Normal']
    s.font.name = 'Garamond'
    s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.15
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

def cp(doc, text, size=11, bold=False, italic=False, sc=False, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if sc: r.font.small_caps = True

def heading(doc, text, level=1):
    sz = {1: 16, 2: 13}.get(level, 11)
    sb = {1: 30, 2: 20}.get(level, 8)
    cp(doc, text.strip(), size=sz, bold=True, sc=(level==1), sb=sb, sa=12)

def body(doc, text):
    t = fix_text(text.strip())
    if not t: return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t)
    r.font.size = Pt(11)

def ast(doc):
    cp(doc, '⁂', size=14, sb=12, sa=12)

def title_page(doc):
    for _ in range(6): doc.add_paragraph('')
    cp(doc, 'NAPOLEON', size=28, bold=True, sc=True)
    cp(doc, 'BY', size=12)
    cp(doc, 'BERTHOLD VALLENTIN', size=16, sc=True)
    for _ in range(2): doc.add_paragraph('')
    cp(doc, 'Translated from the German edition\nBerlin: Georg Bondi, 1923', size=10, italic=True)
    cp(doc, 'Rendered into English in the register\nof E. O. Lorimer\'s 1931 translation of\nKantorowicz, Frederick the Second', size=10, italic=True)
    doc.add_page_break()
    for _ in range(8): doc.add_paragraph('')
    cp(doc, 'HODIERNO HEROI', size=14, italic=True, sc=True)
    cp(doc, '—', size=12, sb=12)
    cp(doc, 'To the Hero of this Day', size=12, sb=6)
    doc.add_page_break()

# ─── Main ─────────────────────────────────────────────────────────────
def build():
    doc = Document()
    setup(doc)
    title_page(doc)
    
    book_page_count = 0
    last_label = None
    
    for fname, label, chap_title in CHAPTERS:
        path = os.path.join(WORKDIR, fname)
        with open(path, 'r', encoding='utf-8') as f:
            lines = [l.rstrip() for l in f.readlines()]
        
        # Handle book-level label (adds page break + book heading)
        if label:
            if label.startswith('BOOK ') or label in ('INTRODUCTION', 'CONCLUSION', 'CHRONOLOGY'):
                book_page_count += 1
                if book_page_count > 1:
                    doc.add_page_break()
                if label != last_label:
                    heading(doc, label, level=1)
                    last_label = label
        
        # Add chapter heading (if not the dedication)
        if chap_title and fname != 'dedication.en.txt':
            # Page break before chapter (unless label already added one)
            if not label:
                doc.add_page_break()
            heading(doc, chap_title, level=2)
        
        # Process content lines — skip the first 1-3 heading lines
        # Pattern: line0=heading, line1=blank, line2=optional title, then text
        start_idx = 0
        # Skip blank lines at beginning
        while start_idx < len(lines) and not lines[start_idx].strip():
            start_idx += 1
        # First non-blank line is the heading — skip it
        start_idx += 1
        # Skip blank line after heading
        while start_idx < len(lines) and not lines[start_idx].strip():
            start_idx += 1
        # If there's a standalone chapter title (like "ACTIVITY" after heading), skip it
        if start_idx < len(lines):
            s = lines[start_idx].strip()
            # These are the known title lines that duplicate info already in chap_title
            known_titles = {
                'ACTIVITY', 'INTENSITY', 'IMMEDIACY: ACTION', 'IMMEDIACY: EXPERIENCE',
                'GOD', 'CONFESSION', 'CHURCH. RITE',
                'MYTH AND DOGMA: YOUTH PERIOD', 'MYTH AND DOGMA: MATURITY',
                'FAITH AND THE STATE',
                'POETRY: DISPOSITIONS AND PRACTICE', 'POETIC INCLINATIONS',
                'TRAGEDY', 'TRAGEDY AND MYTH', 'POETRY AND POLITICS',
                'VISUAL ARTS', 'ARCHITECTURE', 'MUSIC', 'PORTRAITS',
                "ONE'S OWN FEELING: IMMEDIACY",
                'MISSION. SELF-FEELING. POLITICAL SENTIMENT',
                'CURRENTS OF THE AGE',
                'CLASSICAL DISPOSITIONS', 'CLASSICIST PREREQUISITES OF THE AGE',
                'THE CLASSICAL BLOOD', 'THE CLASSICAL SPIRIT',
                'THE CLASSICAL TYPE NAPOLEON: PICTORIAL REPRESENTATION',
                'THE CLASSICAL TYPE NAPOLEON: REPORTS',
                'THE CLASSICAL MAN: END OF THE CLASSICAL TRADITION',
                'THE EXPERIENCE OF HISTORY',
                'TRANSFORMATIONS OF HISTORICAL EXPERIENCE',
                'HISTORY AND DEED',
            }
            if s in known_titles:
                start_idx += 1  # skip it (redundant with chap_title)
                # Skip any blank after title
                while start_idx < len(lines) and not lines[start_idx].strip():
                    start_idx += 1
        
        # Process remaining content
        for i in range(start_idx, len(lines)):
            s = lines[i].strip()
            if not s:
                continue
            if ASTERISM in s:
                ast(doc)
            else:
                body(doc, s)
    
    doc.save(OUTPUT)
    print(f"Done! → {OUTPUT}")
    sz = os.path.getsize(OUTPUT)
    print(f"Output: {sz:,} bytes ({sz/1024:.0f} KB)")

if __name__ == '__main__':
    build()
